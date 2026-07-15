import asyncio
import re
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import RGBColor

from app.core.config import settings
from app.core.paths import storage_key, storage_path
from app.models.employee import Employee
from app.models.order_type import OrderType
from app.schemas.order import OrderCreate
from app.services.template_replacements import (
    build_document_replacements,
    build_order_replacements,
    build_group_order_general_replacements,
    build_group_order_employee_replacements,
    build_vacation_group_extra,
    build_call_group_extra,
    EMPLOYEES_BLOCK_START,
    EMPLOYEES_BLOCK_END,
    APPLICATIONS_BLOCK_START,
    APPLICATIONS_BLOCK_END,
)

MISSING_TEMPLATE_WARNING = "ВНИМАНИЕ: документ сгенерирован без шаблона."


def _replace_text_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
    """Replace placeholders in a paragraph using run coordinate mapping."""
    _replace_in_paragraph(paragraph, replacements)


def _find_block_paragraph_indexes(container: Any, start_marker: str, end_marker: str) -> tuple[int | None, int | None]:
    """Find the paragraph indexes for a block delimited by markers."""
    start_idx = None
    end_idx = None

    for idx, paragraph in enumerate(container.paragraphs):
        text = paragraph.text
        if start_marker in text:
            start_idx = idx
        if end_marker in text and start_idx is not None:
            end_idx = idx
            break

    if start_idx is None or end_idx is None:
        return None, None

    if end_idx <= start_idx:
        raise ValueError(f"Некорректный блок: {start_marker} ... {end_marker}")

    return start_idx, end_idx


def _iter_paragraph_containers(element: Any):
    """Yield element itself and all nested table cells as paragraph containers."""
    yield element
    for table in element.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraph_containers(cell)


def _iter_tables(element: Any):
    """Yield tables in element recursively, including nested tables."""
    for table in element.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_tables(cell)


def _row_contains_marker(row: Any, marker: str) -> bool:
    """Check whether any paragraph in the row contains marker text."""
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            if marker in paragraph.text:
                return True
    return False


def _render_repeat_block_in_table_rows(
    table: Any,
    start_marker: str,
    end_marker: str,
    rows_replacements: list[dict[str, str]],
) -> None:
    """Render repeat blocks where markers are placed in different rows of the same table."""
    from copy import deepcopy
    from docx.table import _Row
    from docx.text.paragraph import Paragraph

    while True:
        start_idx = None
        end_idx = None

        for idx, row in enumerate(table.rows):
            if _row_contains_marker(row, start_marker):
                start_idx = idx
            if _row_contains_marker(row, end_marker) and start_idx is not None:
                end_idx = idx
                break

        if start_idx is None:
            return

        if end_idx is None or end_idx <= start_idx:
            raise ValueError(f"Некорректный блок в таблице: {start_marker} ... {end_marker}")

        template_rows = list(table.rows[start_idx + 1 : end_idx])
        if not template_rows:
            # Special case: marker row may also contain template paragraphs in the same cell.
            start_row = table.rows[start_idx]
            end_row = table.rows[end_idx]

            start_cell = None
            start_paragraph_idx = None
            for cell in start_row.cells:
                for paragraph_idx, paragraph in enumerate(cell.paragraphs):
                    if start_marker in paragraph.text:
                        start_cell = cell
                        start_paragraph_idx = paragraph_idx
                        break
                if start_cell is not None:
                    break

            end_marker_paragraphs = []
            for cell in end_row.cells:
                for paragraph in cell.paragraphs:
                    if end_marker in paragraph.text:
                        end_marker_paragraphs.append(paragraph)

            if start_cell is not None and start_paragraph_idx is not None:
                tail_template_paragraphs = start_cell.paragraphs[start_paragraph_idx + 1 :]
                if tail_template_paragraphs:
                    template_tail_xml = [deepcopy(paragraph._p) for paragraph in tail_template_paragraphs]
                    paragraphs_to_delete = [start_cell.paragraphs[start_paragraph_idx]._p] + [p._p for p in tail_template_paragraphs]

                    for row_replacements in rows_replacements:
                        for paragraph_xml in template_tail_xml:
                            new_paragraph_xml = deepcopy(paragraph_xml)
                            start_cell._tc.append(new_paragraph_xml)
                            _replace_text_in_paragraph(Paragraph(new_paragraph_xml, start_cell), row_replacements)

                    for paragraph_xml in paragraphs_to_delete:
                        paragraph_xml.getparent().remove(paragraph_xml)

                    for paragraph in end_marker_paragraphs:
                        paragraph._p.getparent().remove(paragraph._p)

                    # If end row became empty, remove it.
                    row_has_text = any(
                        paragraph.text.strip()
                        for cell in end_row.cells
                        for paragraph in cell.paragraphs
                    )
                    if not row_has_text:
                        end_row._tr.getparent().remove(end_row._tr)
                    continue

            # Nothing between markers — clear marker rows
            for row in (table.rows[start_idx], table.rows[end_idx]):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = paragraph.text.replace(start_marker, "").replace(end_marker, "")
            continue

        template_row_xml = [deepcopy(row._tr) for row in template_rows]
        rows_to_delete = [table.rows[start_idx]._tr, *[row._tr for row in template_rows], table.rows[end_idx]._tr]

        insert_before = table.rows[end_idx]._tr
        for row_replacements in rows_replacements:
            for row_xml in template_row_xml:
                new_row_xml = deepcopy(row_xml)
                insert_before.addprevious(new_row_xml)
                inserted_row = _Row(new_row_xml, table)
                for cell in inserted_row.cells:
                    _replace_placeholders_in_element(cell, row_replacements)

        for row_elem in rows_to_delete:
            row_elem.getparent().remove(row_elem)


def _render_body_level_block_between_markers(
    doc: Document,
    start_paragraph: Any,
    end_paragraph: Any,
    rows_replacements: list[dict[str, str]],
) -> bool:
    """Render a block where markers are body paragraphs and template content between them may include tables."""
    from copy import deepcopy
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body_element = doc._body._element
    start_elem = start_paragraph._element
    end_elem = end_paragraph._element

    if start_elem.getparent() is not body_element or end_elem.getparent() is not body_element:
        return False

    body_children = list(body_element)
    try:
        start_child_idx = body_children.index(start_elem)
        end_child_idx = body_children.index(end_elem)
    except ValueError:
        return False

    if end_child_idx <= start_child_idx:
        return False

    template_elements = body_children[start_child_idx + 1 : end_child_idx]
    if not template_elements:
        return False

    insert_before = end_paragraph._p
    parent = end_paragraph._parent

    for row_replacements in rows_replacements:
        for template_elem in template_elements:
            new_elem = deepcopy(template_elem)
            insert_before.addprevious(new_elem)

            tag = new_elem.tag.rsplit("}", 1)[-1]
            if tag == "p":
                _replace_text_in_paragraph(Paragraph(new_elem, parent), row_replacements)
            elif tag == "tbl":
                _replace_placeholders_in_element(Table(new_elem, parent), row_replacements)

    elements_to_delete = [start_elem, *template_elements, end_elem]
    for elem in elements_to_delete:
        elem.getparent().remove(elem)

    return True


def _render_repeat_block(
    doc: Document,
    start_marker: str,
    end_marker: str,
    rows_replacements: list[dict[str, str]],
) -> None:
    """Find ALL blocks between markers in doc and table cells, replicate each, and substitute placeholders."""
    from copy import deepcopy
    from docx.text.paragraph import Paragraph

    def _render_repeat_block_in_container(container: Any) -> None:
        while True:
            start_idx, end_idx = _find_block_paragraph_indexes(container, start_marker, end_marker)

            if start_idx is None:
                return

            paragraphs = container.paragraphs
            start_p = paragraphs[start_idx]
            end_p = paragraphs[end_idx]

            # Paragraphs inside the block that should be replicated
            template_paragraphs = paragraphs[start_idx + 1:end_idx]

            if not template_paragraphs:
                if _render_body_level_block_between_markers(doc, start_p, end_p, rows_replacements):
                    continue
                # Nothing between markers — just clear the markers
                start_p.text = ""
                end_p.text = ""
                continue

            # Save XML copies of template paragraphs
            template_xml = [deepcopy(p._p) for p in template_paragraphs]

            # Collect all paragraph elements that need to be deleted (markers + template)
            paragraphs_to_delete = [paragraphs[start_idx]._element]  # start marker
            for p in template_paragraphs:
                paragraphs_to_delete.append(p._element)
            paragraphs_to_delete.append(paragraphs[end_idx]._element)  # end marker

            # Insert new paragraphs before end_marker, one set per employee
            insert_before = end_p._p

            for row_replacements in rows_replacements:
                for p_xml in template_xml:
                    new_p_xml = deepcopy(p_xml)
                    insert_before.addprevious(new_p_xml)
                    # Replace placeholders immediately on the inserted paragraph.
                    # This avoids relying on element identity checks across lxml proxies.
                    inserted_paragraph = Paragraph(new_p_xml, end_p._parent)
                    _replace_text_in_paragraph(inserted_paragraph, row_replacements)

            # Remove template paragraphs and markers (use collected elements, not indices)
            for p_elem in paragraphs_to_delete:
                p_elem.getparent().remove(p_elem)

    for container in _iter_paragraph_containers(doc):
        _render_repeat_block_in_container(container)

    for table in _iter_tables(doc):
        _render_repeat_block_in_table_rows(table, start_marker, end_marker, rows_replacements)


def _replace_placeholders_in_element(element: Any, replacements: dict[str, str]) -> None:
    """Replace placeholders in a document element (paragraphs and tables)."""
    if hasattr(element, "paragraphs"):
        for paragraph in element.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)

    if hasattr(element, "tables"):
        for table in element.tables:
            _replace_placeholders_in_element(table, replacements)

    if hasattr(element, "rows") and hasattr(element, "columns"):
        for row in element.rows:
            for cell in row.cells:
                _replace_placeholders_in_element(cell, replacements)


async def generate_document(
    order_number: str,
    data: OrderCreate,
    employee: Employee,
    order_type: OrderType,
    year_dir: Path,
) -> tuple[str, str]:
    """Generate a DOCX document for a single-employee order."""
    doc, replacements = await _build_document(order_number, data, employee, order_type)

    storage_name = _build_storage_name(
        order_number, data.order_date, order_type, employee,
        _extract_extra_dates(data.extra_fields, order_type.code),
    )
    display_name = _build_display_name(
        order_number, data.order_date, order_type, employee,
        _extract_extra_info(data.extra_fields, order_type.code),
    )
    file_path = year_dir / storage_name

    if data.draft_id:
        # Copy only — draft is deleted after successful order create (order_service).
        from app.services.order_draft_service import order_draft_service

        draft_path = order_draft_service.get_draft_path(data.draft_id)
        file_path.parent.mkdir(parents=True, exist_ok=True)
        import shutil
        shutil.copy2(str(draft_path), str(file_path))
        return storage_key(file_path, "ORDERS_PATH"), display_name

    await asyncio.wait_for(
        asyncio.to_thread(doc.save, str(file_path)),
        timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
    )
    return storage_key(file_path, "ORDERS_PATH"), display_name


def _build_group_general_replacements(
    employee_rows: list[dict],
    order_number: str,
    order_date: date,
    extra: dict[str, str],
) -> dict[str, str]:
    """Build general (non-per-employee) replacements for group orders."""
    from app.services.template_replacements import build_template_replacements_for_employee

    first_emp = employee_rows[0]["employee"] if employee_rows else None
    emp_replacements = build_template_replacements_for_employee(first_emp)

    return {
        "{order_number}": order_number,
        "{order_date}": order_date.strftime("%d.%m.%Y"),
        "{oznak_gender}": emp_replacements.get("{oznak}", ""),
        "{initials_before}": emp_replacements.get("{initials_before}", ""),
        **extra,
    }


def _build_employee_replacements(idx: int, emp: Employee, extra: dict[str, str]) -> dict[str, str]:
    """Build per-employee placeholder replacements, consistent with _prepare_replacements for single orders."""
    from app.services.template_replacements import build_template_replacements_for_employee

    return {
        "{index}": str(idx),
        **build_template_replacements_for_employee(emp),
        "{position}": str(emp.position.name if emp.position else "").lower(),
        **extra,
    }


async def render_vacation_unpaid_group_docx(
    order_number: str,
    data: "VacationUnpaidGroupOrderCreate",
    order_type: OrderType,
    employee_rows: list[dict],
    output_path: Path | None = None,
) -> bytes | None:
    """
    Render group vacation unpaid order DOCX from template.

    If output_path is provided, save the DOCX there and return None.
    If output_path is not provided, return DOCX bytes.
    """
    template_name = order_type.template_filename or "template__order__vacation_unpaid_group.docx"
    template_path = get_template_path_for_filename(template_name)

    if template_path.exists():
        doc = await asyncio.wait_for(
            asyncio.to_thread(Document, str(template_path)),
            timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
        )
    else:
        doc = Document()
        doc.add_heading(f"Приказ №{order_number}", level=1)
        warning_run = doc.add_paragraph().add_run(MISSING_TEMPLATE_WARNING)
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    # Build employee block replacements
    employee_rows_replacements = []
    for idx, row_data in enumerate(employee_rows, 1):
        emp = row_data["employee"]
        vacation_days = row_data["vacation_days"]

        employee_rows_replacements.append(build_group_order_employee_replacements(idx, emp, build_vacation_group_extra(
            vacation_start=data.vacation_start,
            vacation_end=row_data["vacation_end"],
            vacation_days=vacation_days,
            application_date=row_data.get("application_date", data.order_date),
        )))

    # Render employee blocks
    _render_repeat_block(
        doc,
        EMPLOYEES_BLOCK_START,
        EMPLOYEES_BLOCK_END,
        employee_rows_replacements,
    )

    # Render application blocks
    _render_repeat_block(
        doc,
        APPLICATIONS_BLOCK_START,
        APPLICATIONS_BLOCK_END,
        employee_rows_replacements,
    )

    # General placeholders
    replacements = build_group_order_general_replacements(
        employee_rows,
        order_number,
        data.order_date,
        extra={
            "{vacation_start}": data.vacation_start.strftime("%d.%m.%Y"),
            "{vacation_end}": data.vacation_start.strftime("%d.%m.%Y"),  # fallback; per-employee values override in blocks
        },
    )

    _replace_placeholders_in_element(doc, replacements)

    if output_path is not None:
        await asyncio.wait_for(
            asyncio.to_thread(doc.save, str(output_path)),
            timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
        )
        return None

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def generate_group_document(
    order_number: str,
    data: "VacationUnpaidGroupOrderCreate",
    order_type: OrderType,
    year_dir: Path,
    employee_rows: list[dict],
) -> tuple[str, str]:
    """Generate DOCX for a group order using the group template."""
    storage_name = f"prikaz_{order_number}_vacation_unpaid_group_{data.vacation_start.strftime('%Y-%m-%d')}.docx"
    display_name = f"Приказ №{order_number} от {data.order_date.strftime('%d.%m.%Y')} — отпуск за свой счет (групповой, {len(employee_rows)} сотр.).docx"
    file_path = year_dir / storage_name

    await render_vacation_unpaid_group_docx(
        order_number, data, order_type, employee_rows, output_path=file_path,
    )

    return storage_key(file_path, "ORDERS_PATH"), display_name


async def generate_weekend_call_group_document(
    order_number: str,
    data: "WeekendCallGroupOrderCreate",
    order_type: OrderType,
    year_dir: Path,
    employee_rows: list[dict],
    call_start: date,
    call_end: date,
    output_path: Path | None = None,
) -> tuple[str, str]:
    """Generate DOCX for a group weekend call order using the group template."""
    template_name = order_type.template_filename or "template__order__weekend_call_group.docx"
    template_path = get_template_path_for_filename(template_name)

    if template_path.exists():
        doc = await asyncio.wait_for(
            asyncio.to_thread(Document, str(template_path)),
            timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
        )
    else:
        doc = Document()
        doc.add_heading(f"Приказ №{order_number}", level=1)
        warning_run = doc.add_paragraph().add_run(MISSING_TEMPLATE_WARNING)
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    # Build employee block replacements
    employee_rows_replacements = []
    for idx, row_data in enumerate(employee_rows, 1):
        emp = row_data["employee"]
        vacation_days = row_data["vacation_days"]

        employee_rows_replacements.append(build_group_order_employee_replacements(idx, emp, build_call_group_extra(
            call_start=call_start,
            call_end=call_end,
            vacation_days=vacation_days,
            application_date=row_data.get("application_date", data.order_date),
        )))

    # Render employee blocks
    _render_repeat_block(
        doc,
        EMPLOYEES_BLOCK_START,
        EMPLOYEES_BLOCK_END,
        employee_rows_replacements,
    )

    # Render application blocks
    _render_repeat_block(
        doc,
        APPLICATIONS_BLOCK_START,
        APPLICATIONS_BLOCK_END,
        employee_rows_replacements,
    )

    # General placeholders
    # Format call_date: single date if start==end, otherwise range
    if call_start == call_end:
        call_date_display = call_start.strftime("%d.%m.%Y")
    else:
        call_date_display = f"{call_start.strftime('%d.%m.%Y')} по {call_end.strftime('%d.%m.%Y')}"

    replacements = build_group_order_general_replacements(
        employee_rows,
        order_number,
        data.order_date,
        extra={
            "{call_date}": call_date_display,
            "{call_date_start}": call_start.strftime("%d.%m.%Y"),
            "{call_date_end}": call_end.strftime("%d.%m.%Y"),
        },
    )

    _replace_placeholders_in_element(doc, replacements)

    if output_path:
        file_path = output_path
        storage_name = output_path.name
    else:
        storage_name = f"prikaz_{order_number}_weekend_call_group_{call_start.strftime('%Y-%m-%d')}.docx"
        file_path = year_dir / storage_name

    display_name = f"Приказ №{order_number} от {data.order_date.strftime('%d.%m.%Y')} — вызов в выходной (групповой, {len(employee_rows)} сотр.).docx"

    await asyncio.wait_for(
        asyncio.to_thread(doc.save, str(file_path)),
        timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
    )
    
    if output_path:
        return output_path.name, display_name
    return storage_key(file_path, "ORDERS_PATH"), display_name


def copy_docx_to_permanent(source_path: Path, destination_path: Path) -> None:
    """
    Copy a DOCX file to the permanent order file path.
    This function does not know about draft_id.
    """
    destination_path.parent.mkdir(parents=True, exist_ok=True)
    import shutil
    shutil.copy2(str(source_path), str(destination_path))


def get_template_path(order_type: OrderType) -> Path:
    """Get the template file path for an order type."""
    if not order_type.template_filename:
        return Path(settings.TEMPLATES_PATH) / "__missing__.docx"
    return storage_path(order_type.template_filename, "TEMPLATES_PATH")


def get_template_path_for_filename(filename: str) -> Path:
    """Get the template file path for a filename."""
    return Path(settings.TEMPLATES_PATH) / filename


async def _build_document(
    order_number: str,
    data: OrderCreate,
    employee: Employee | None,
    order_type: OrderType,
) -> tuple[Document, dict[str, str]]:
    template_path = get_template_path(order_type)

    if template_path.exists():
        doc = await asyncio.wait_for(
            asyncio.to_thread(Document, str(template_path)),
            timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
        )
    else:
        doc = Document()
        doc.add_heading(f"Приказ №{order_number}", level=1)
        warning_run = doc.add_paragraph().add_run(MISSING_TEMPLATE_WARNING)
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        doc.add_paragraph(f"Тип: {order_type.name}")
        doc.add_paragraph(f"Дата: {data.order_date.strftime('%d.%m.%Y')}")
        if employee:
            doc.add_paragraph(f"Сотрудник: {employee.name}")

    replacements = build_order_replacements(
        order_number=order_number,
        order_date=data.order_date,
        order_type_name=order_type.name,
        order_type_code=order_type.code,
        employee=employee,
        extra_fields=data.extra_fields,
        notes=data.notes or "",
    )
    await asyncio.wait_for(
        asyncio.to_thread(_replace_placeholders, doc, replacements),
        timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
    )
    return doc, replacements


def _replace_placeholders(target: Any, replacements: dict[str, str]) -> None:
    """Replace placeholders in a Document, Paragraph, or Cell. Dispatches to the appropriate handler."""
    if hasattr(target, "paragraphs") and hasattr(target, "tables"):
        # Document
        for paragraph in target.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in target.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)
    elif hasattr(target, "paragraphs"):
        # Cell or similar object with paragraphs
        for paragraph in target.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
    else:
        # Single Paragraph
        _replace_in_paragraph(target, replacements)


def _replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
    """Replace placeholders in a single paragraph using run coordinate mapping.
    All replacements are applied in a single pass (right-to-left) so run indices stay valid."""
    if not paragraph.runs:
        return

    full_text = paragraph.text

    # Collect all occurrences of all keys
    occurrences: list[tuple[int, str, str]] = []  # (start_pos, key, value)
    for key, value in replacements.items():
        if key not in full_text:
            continue
        key_len = len(key)
        for i in range(len(full_text)):
            if full_text.startswith(key, i):
                occurrences.append((i, key, value))

    if not occurrences:
        return

    # Sort by position descending (right-to-left) so earlier indices stay valid
    occurrences.sort(key=lambda x: x[0], reverse=True)

    # Build coordinate map once
    p_map = []
    for run_idx, run in enumerate(paragraph.runs):
        for char_idx, char in enumerate(run.text):
            p_map.append({"run": run_idx, "char": char_idx})

    for start_pos, key, value in occurrences:
        key_len = len(key)
        # If the key spans more characters than we have in the map, skip (likely split across runs in a way we can't handle)
        if start_pos + key_len > len(p_map):
            continue
        key_map = p_map[start_pos : start_pos + key_len]
        _replace_in_runs(paragraph.runs, key_map, value)


def _replace_in_runs(runs: list[Any], key_map: list[dict], value: str) -> None:
    """Apply replacement to the specific runs/characters identified by key_map."""
    for i, position in enumerate(reversed(key_map), start=1):
        run_idx = position["run"]
        char_idx = position["char"]
        run = runs[run_idx]
        chars = list(run.text)

        if i < len(key_map):
            # Not the first character of the key — delete it
            chars.pop(char_idx)
        else:
            # First character (last in reversed order) — replace with value
            chars[char_idx] = value

        run.text = "".join(chars)


def _extract_extra_dates(extra_fields: dict | None, order_type_code: str) -> list[date] | None:
    """Extract relevant dates from extra_fields for storage filename."""
    if not extra_fields:
        return None
    dates: list[date] = []
    try:
        if order_type_code in ("vacation_paid", "vacation_unpaid"):
            if extra_fields.get("vacation_start"):
                dates.append(date.fromisoformat(extra_fields["vacation_start"]))
            if extra_fields.get("vacation_end"):
                dates.append(date.fromisoformat(extra_fields["vacation_end"]))
        elif order_type_code == "vacation_recall":
            if extra_fields.get("recall_date"):
                dates.append(date.fromisoformat(extra_fields["recall_date"]))
        elif order_type_code == "vacation_postpone":
            if extra_fields.get("new_vacation_start"):
                dates.append(date.fromisoformat(extra_fields["new_vacation_start"]))
            if extra_fields.get("new_vacation_end"):
                dates.append(date.fromisoformat(extra_fields["new_vacation_end"]))
        elif order_type_code == "vacation_extension":
            if extra_fields.get("sick_start_date"):
                dates.append(date.fromisoformat(extra_fields["sick_start_date"]))
            if extra_fields.get("sick_end_date"):
                dates.append(date.fromisoformat(extra_fields["sick_end_date"]))
    except (ValueError, TypeError):
        pass
    return dates or None


def _extract_extra_info(extra_fields: dict | None, order_type_code: str) -> str:
    """Extract human-readable extra info for display name."""
    if not extra_fields:
        return ""
    parts: list[str] = []
    if order_type_code in ("vacation_paid", "vacation_unpaid"):
        start = extra_fields.get("vacation_start", "")
        end = extra_fields.get("vacation_end", "")
        if start and end:
            try:
                s = date.fromisoformat(start).strftime("%d.%m.%Y")
                e = date.fromisoformat(end).strftime("%d.%m.%Y")
                parts.append(f"{s}-{e}")
            except ValueError:
                pass
    elif order_type_code == "vacation_recall":
        recall = extra_fields.get("recall_date", "")
        if recall:
            try:
                parts.append(f"с {date.fromisoformat(recall).strftime('%d.%m.%Y')}")
            except ValueError:
                pass
    elif order_type_code == "vacation_postpone":
        ns = extra_fields.get("new_vacation_start", "")
        ne = extra_fields.get("new_vacation_end", "")
        if ns and ne:
            try:
                parts.append(f"{date.fromisoformat(ns).strftime('%d.%m.%Y')}-{date.fromisoformat(ne).strftime('%d.%m.%Y')}")
            except ValueError:
                pass
    elif order_type_code == "vacation_extension":
        ss = extra_fields.get("sick_start_date", "")
        se = extra_fields.get("sick_end_date", "")
        if ss and se:
            try:
                parts.append(f"больничный {date.fromisoformat(ss).strftime('%d.%m.%Y')}-{date.fromisoformat(se).strftime('%d.%m.%Y')}")
            except ValueError:
                pass
    return " ".join(parts)


def _transliterate(text: str) -> str:
    """Convert Cyrillic text to Latin transcription."""
    mapping = {
        'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'yo',
        'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
        'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
        'ф': 'f', 'х': 'kh', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'shch',
        'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
        'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'Yo',
        'Ж': 'Zh', 'З': 'Z', 'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M',
        'Н': 'N', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U',
        'Ф': 'F', 'Х': 'Kh', 'Ц': 'Ts', 'Ч': 'Ch', 'Ш': 'Sh', 'Щ': 'Shch',
        'Ъ': '', 'Ы': 'Y', 'Ь': '', 'Э': 'E', 'Ю': 'Yu', 'Я': 'Ya',
    }
    result = []
    for ch in text:
        result.append(mapping.get(ch, ch))
    return ''.join(result)


def _build_storage_name(
    order_number: str,
    order_date: date,
    order_type: OrderType,
    employee: Employee | None,
    extra_dates: list[date] | None = None,
) -> str:
    """Build a filesystem-safe filename (ASCII only, no spaces)."""
    type_code = order_type.code
    if employee and employee.name:
        last_name = employee.name.split()[0]
        transliterated = _transliterate(last_name).lower()
        transliterated = re.sub(r'[^a-z0-9-]', '', transliterated.replace(' ', '-'))
    else:
        transliterated = "group"

    parts = [
        order_date.isoformat(),
        f"prikaz_{order_number}",
        type_code,
        transliterated,
    ]
    if extra_dates:
        date_strs = [d.isoformat() for d in extra_dates]
        parts.append("_".join(date_strs))
    return "_".join(parts) + ".docx"


def _build_display_name(
    order_number: str,
    order_date: date,
    order_type: OrderType,
    employee: Employee | None,
    extra_info: str = "",
) -> str:
    """Build a human-readable display name in Russian."""
    date_str = order_date.strftime("%d.%m.%Y")
    emp_name = employee.name if employee else "Групповой"
    name = f"Приказ №{order_number} от {date_str} - {order_type.name} - {emp_name}"
    if extra_info:
        name += f" - {extra_info}"
    return name + ".docx"


def _build_filename(order_number: str, order_type: OrderType, replacements: dict[str, str]) -> str:
    pattern = order_type.filename_pattern or "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx"
    filename = pattern
    for key, value in replacements.items():
        filename = filename.replace(key, value)
    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"
    sanitized = re.sub(r'[<>:"/\\\\|?*]+', "_", filename).strip()
    return sanitized or f"order_{order_number}.docx"
