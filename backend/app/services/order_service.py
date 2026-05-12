import asyncio
import re
import shutil
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import RGBColor
from sqlalchemy import delete as sa_delete, or_, select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.exceptions import DuplicateError, EmployeeNotFoundError, HRMSException, OrderNotFoundError, VacationOverlapError
from app.core.logging import get_audit_logger
from app.core.paths import storage_key, storage_path
from app.models.employee import Employee
from app.models.order import Order
from app.models.order_type import OrderType
from app.repositories.employee_repository import EmployeeRepository
from app.repositories.order_repository import OrderRepository
from app.repositories.order_type_repository import OrderTypeRepository
from app.schemas.order import OrderCreate, OrderUpdate
from app.schemas.order_type import OrderTypeCreate, OrderTypeUpdate

audit_logger = get_audit_logger()
MISSING_TEMPLATE_WARNING = "ВНИМАНИЕ: документ сгенерирован без шаблона."


DEFAULT_ORDER_TYPES: list[dict[str, Any]] = [
    {
        "code": "hire",
        "name": "Прием на работу",
        "show_in_orders_page": True,
        "template_filename": "prikaz_priem.docx",
        "field_schema": [
            {"key": "hire_date", "label": "Дата приема", "type": "date", "required": False},
            {"key": "contract_end", "label": "Конец контракта", "type": "date", "required": False},
            {"key": "trial_end", "label": "Конец испытательного срока", "type": "date", "required": False},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "к",
    },
    {
        "code": "dismissal",
        "name": "Увольнение",
        "show_in_orders_page": True,
        "template_filename": "prikaz_uvolnenie.docx",
        "field_schema": [
            {"key": "dismissal_date", "label": "Дата увольнения", "type": "date", "required": False},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "к",
    },
    {
        "code": "transfer",
        "name": "Перевод",
        "show_in_orders_page": True,
        "template_filename": "prikaz_perevod.docx",
        "field_schema": [
            {"key": "transfer_date", "label": "Дата перевода", "type": "date", "required": False},
            {"key": "transfer_reason", "label": "Основание", "type": "textarea", "required": False},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "к",
    },
    {
        "code": "contract_extension",
        "name": "Продление контракта",
        "show_in_orders_page": True,
        "template_filename": "prikaz_prodlenie_kontrakta.docx",
        "field_schema": [
            {"key": "contract_new_end", "label": "Новая дата конца контракта", "type": "date", "required": False},
            {"key": "trial_end", "label": "Конец испытательного срока", "type": "date", "required": False},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "к",
    },
    {
        "code": "vacation_paid",
        "name": "Отпуск трудовой",
        "show_in_orders_page": False,
        "template_filename": "prikaz_otpusk_trudovoy.docx",
        "field_schema": [
            {"key": "vacation_start", "label": "Дата начала", "type": "date", "required": True},
            {"key": "vacation_end", "label": "Дата окончания", "type": "date", "required": True},
            {"key": "vacation_days", "label": "Количество дней", "type": "number", "required": True},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "л",
    },
    {
        "code": "vacation_unpaid",
        "name": "Отпуск за свой счет",
        "show_in_orders_page": False,
        "template_filename": "prikaz_otpusk_svoy_schet.docx",
        "field_schema": [
            {"key": "vacation_start", "label": "Дата начала", "type": "date", "required": True},
            {"key": "vacation_end", "label": "Дата окончания", "type": "date", "required": True},
            {"key": "vacation_days", "label": "Количество дней", "type": "number", "required": True},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "к",
    },
    {
        "code": "weekend_call",
        "name": "Вызов в выходной",
        "show_in_orders_page": False,
        "template_filename": "prikaz_vyzov_v_vyhodnoy.docx",
        "field_schema": [
            {"key": "call_date", "label": "Дата вызова", "type": "date", "required": False},
            {"key": "call_date_start", "label": "Дата начала", "type": "date", "required": False},
            {"key": "call_date_end", "label": "Дата окончания", "type": "date", "required": False},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "к",
    },
    {
        "code": "vacation_recall",
        "name": "Отзыв из отпуска",
        "show_in_orders_page": False,
        "template_filename": "prikaz_otzyv_iz_otpuska.docx",
        "field_schema": [
            {"key": "recall_date", "label": "Дата отзыва", "type": "date", "required": True},
            {"key": "old_vacation_start", "label": "Дата начала отпуска", "type": "date", "required": True},
            {"key": "old_vacation_end", "label": "Дата окончания отпуска", "type": "date", "required": True},
            {"key": "old_vacation_days", "label": "Количество дней отпуска", "type": "number", "required": True},
            {"key": "reason", "label": "Основание", "type": "text", "required": False},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "л",
    },
    {
        "code": "vacation_postpone",
        "name": "Перенос отпуска",
        "show_in_orders_page": False,
        "template_filename": "prikaz_perenos_otpuska.docx",
        "field_schema": [
            {"key": "old_vacation_start", "label": "Старая дата начала", "type": "date", "required": True},
            {"key": "old_vacation_end", "label": "Старая дата окончания", "type": "date", "required": True},
            {"key": "new_vacation_start", "label": "Новая дата начала", "type": "date", "required": True},
            {"key": "new_vacation_end", "label": "Новая дата окончания", "type": "date", "required": True},
            {"key": "vacation_days", "label": "Количество дней", "type": "number", "required": True},
            {"key": "reason", "label": "Основание", "type": "text", "required": False},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "л",
    },
    {
        "code": "vacation_extension",
        "name": "Продление отпуска",
        "show_in_orders_page": False,
        "template_filename": "prikaz_prodlenie_otpuska.docx",
        "field_schema": [
            {"key": "vacation_start", "label": "Дата начала отпуска", "type": "date", "required": True},
            {"key": "vacation_end", "label": "Дата окончания отпуска", "type": "date", "required": True},
            {"key": "vacation_days", "label": "Количество дней отпуска", "type": "number", "required": True},
            {"key": "sick_start_date", "label": "Дата начала больничного", "type": "date", "required": True},
            {"key": "sick_end_date", "label": "Дата окончания больничного", "type": "date", "required": True},
            {"key": "comment", "label": "Комментарий", "type": "text", "required": False},
        ],
        "filename_pattern": "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx",
        "letter": "л",
    },
]

STANDARD_ORDER_CODES = frozenset(item["code"] for item in DEFAULT_ORDER_TYPES)


class OrderService:
    def __init__(self):
        self.order_repo = OrderRepository()
        self.order_type_repo = OrderTypeRepository()
        self.employee_repo = EmployeeRepository()

    async def ensure_default_order_types(self, db: AsyncSession) -> list[OrderType]:
        existing = await self.order_type_repo.list_all(db)
        existing_by_code = {item.code: item for item in existing}
        changed = False

        for item in DEFAULT_ORDER_TYPES:
            current = existing_by_code.get(item["code"])
            display_name = f"Шаблон - {item['name']}.docx"
            if not current:
                created = await self.order_type_repo.create(db, {**item, "display_name": display_name})
                existing_by_code[created.code] = created
                changed = True
                continue

            updates: dict[str, Any] = {}
            for key in ("name", "show_in_orders_page", "filename_pattern", "letter"):
                if getattr(current, key) != item.get(key):
                    updates[key] = item.get(key)
            if not current.display_name or current.display_name.startswith("Шаблон - "):
                updates["display_name"] = display_name
            if updates:
                await self.order_type_repo.update(db, current, updates)
                changed = True

        if changed:
            await db.commit()

        return list(existing_by_code.values())

    async def get_order_types(
        self,
        db: AsyncSession,
        active_only: bool = True,
        show_in_orders_page: bool | None = None,
    ) -> list[dict[str, Any]]:
        await self.ensure_default_order_types(db)
        items = await self.order_type_repo.list_all(
            db,
            active_only=active_only,
            show_in_orders_page=show_in_orders_page,
        )
        return [self._serialize_order_type(item) for item in items]

    async def get_order_type(self, db: AsyncSession, order_type_id: int) -> OrderType:
        await self.ensure_default_order_types(db)
        order_type = await self.order_type_repo.get_by_id(db, order_type_id)
        if not order_type:
            raise HRMSException("Тип приказа не найден", "order_type_not_found", status_code=404)
        return order_type

    async def get_order_type_by_code(self, db: AsyncSession, code: str) -> OrderType:
        await self.ensure_default_order_types(db)
        order_type = await self.order_type_repo.get_by_code(db, code)
        if not order_type:
            raise HRMSException("Тип приказа не найден", "order_type_not_found", status_code=404)
        return order_type

    async def create_order_type(self, db: AsyncSession, data: OrderTypeCreate) -> dict[str, Any]:
        await self.ensure_default_order_types(db)
        if await self.order_type_repo.get_by_code(db, data.code):
            raise DuplicateError(f"Тип приказа с кодом {data.code} уже существует", "duplicate_order_type_code")
        if await self.order_type_repo.get_by_name(db, data.name):
            raise DuplicateError(f"Тип приказа с названием {data.name} уже существует", "duplicate_order_type_name")

        created = await self.order_type_repo.create(db, data.model_dump())
        await db.commit()
        return self._serialize_order_type(created)

    async def update_order_type(self, db: AsyncSession, order_type_id: int, data: OrderTypeUpdate) -> dict[str, Any]:
        order_type = await self.get_order_type(db, order_type_id)
        payload = data.model_dump(exclude_unset=True)

        # Стандартные типы приказов нельзя редактировать — их поля задаются миграцией/кодом.
        # Разрешена только загрузка/удаление шаблона.
        if order_type.code in STANDARD_ORDER_CODES:
            blocked = set(payload.keys()) - {"template_filename"}
            if blocked:
                raise HRMSException(
                    f"Нельзя изменить стандартный тип приказа. Заблокированные поля: {', '.join(sorted(blocked))}",
                    "standard_order_type_readonly",
                    status_code=403,
                )
            # Для стандартных типов разрешаем только template_filename (через upload_template)
            # Если пришёл пустой payload — просто возвращаем текущее состояние
            if not payload:
                return self._serialize_order_type(order_type)

        new_name = payload.get("name")
        if new_name and new_name != order_type.name:
            existing = await self.order_type_repo.get_by_name(db, new_name)
            if existing and existing.id != order_type.id:
                raise DuplicateError(f"Тип приказа с названием {new_name} уже существует", "duplicate_order_type_name")

        updated = await self.order_type_repo.update(db, order_type, payload)
        await db.commit()
        return self._serialize_order_type(updated)

    async def delete_order_type(self, db: AsyncSession, order_type_id: int) -> None:
        order_type = await self.get_order_type(db, order_type_id)
        order_count = await self.order_type_repo.count_orders(db, order_type_id)
        if order_count > 0:
            recent_orders = await self.order_repo.get_recent_by_order_type(db, order_type_id, limit=5)
            numbers = [o.order_number for o in recent_orders]
            raise HRMSException(
                f"Нельзя удалить тип приказа, который уже используется. Связанные приказы: {', '.join(numbers)}",
                "order_type_in_use",
                status_code=409,
            )

        template_path = self._get_template_path(order_type)
        if template_path.exists():
            template_path.unlink()

        await self.order_type_repo.delete(db, order_type)
        await db.commit()

    async def get_next_number(self, db: AsyncSession, order_type_id: int) -> str:
        return await self.order_repo.get_next_order_number(db, order_type_id)

    async def get_years(self, db: AsyncSession) -> list[int]:
        return await self.order_repo.get_years(db)

    async def get_all(
        self,
        db: AsyncSession,
        page: int = 1,
        per_page: int = 20,
        sort_by: Optional[str] = None,
        sort_order: str = "desc",
        year: Optional[int] = None,
        order_type_code: Optional[str] = None,
        order_letter: Optional[str] = None,
        employee_id: Optional[int] = None,
        date_from: Optional[date] = None,
        date_to: Optional[date] = None,
        order_number: Optional[str] = None,
    ) -> dict[str, Any]:
        items, total = await self.order_repo.get_all(
            db,
            page=page,
            per_page=per_page,
            sort_by=sort_by,
            sort_order=sort_order,
            year=year,
            order_type_code=order_type_code,
            order_letter=order_letter,
            employee_id=employee_id,
            date_from=date_from,
            date_to=date_to,
            order_number=order_number,
        )
        total_pages = max(1, (total + per_page - 1) // per_page)
        return {
            "items": [self._serialize_order(order) for order in items],
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": total_pages,
        }

    async def get_recent(self, db: AsyncSession, limit: int = 10, year: Optional[int] = None) -> list[dict[str, Any]]:
        items = await self.order_repo.get_recent(db, limit=limit, year=year)
        return [self._serialize_order(order) for order in items]

    async def get_by_id(self, db: AsyncSession, order_id: int) -> Order:
        order = await self.order_repo.get_by_id(db, order_id)
        if not order:
            raise OrderNotFoundError(order_id)
        return order

    async def create_order(self, db: AsyncSession, data: OrderCreate) -> Order:
        await self.ensure_default_order_types(db)
        if db.in_transaction():
            return await self._do_create_order(db, data)
        async with db.begin():
            return await self._do_create_order(db, data)

    async def _do_create_order(self, db: AsyncSession, data: OrderCreate) -> Order:
        order_number = data.order_number
        if not order_number:
            order_number = await self.order_repo.get_next_order_number(db, data.order_type_id)
        else:
            order_number = order_number.strip()

        employee = await self.employee_repo.get_by_id(db, data.employee_id)
        if not employee:
            raise EmployeeNotFoundError(data.employee_id)

        if not data.order_type_id:
            raise HRMSException("Не передан order_type_id", "order_type_not_found", status_code=422)

        order_type = await self.order_type_repo.get_by_id(db, data.order_type_id)
        if not order_type or not order_type.is_active:
            raise HRMSException("Активный тип приказа не найден", "order_type_not_found", status_code=404)

        year_dir = Path(settings.ORDERS_PATH) / str(data.order_date.year)
        year_dir.mkdir(parents=True, exist_ok=True)

        file_path, display_name = await self._generate_document(order_number, data, employee, order_type, year_dir)

        # Подготавливаем extra_fields для продления контракта
        extra_fields = data.extra_fields
        if order_type.code == "contract_extension" and data.extra_fields:
            new_end = data.extra_fields.get("contract_new_end")
            if new_end:
                extra_fields = dict(data.extra_fields)
                extra_fields["old_contract_end"] = (
                    employee.contract_end.isoformat() if employee.contract_end else None
                )
                employee.contract_end = date.fromisoformat(new_end)
                await db.flush()

        order = await self.order_repo.create(
            db,
            {
                "order_number": order_number,
                "order_type_id": order_type.id,
                "employee_id": data.employee_id,
                "order_date": data.order_date,
                "file_path": file_path,
                "display_name": display_name,
                "notes": data.notes,
                "extra_fields": extra_fields,
            },
        )

        # Автоматическая архивация сотрудника при приказе об увольнении
        if order_type.code == "dismissal" and not employee.is_dismissed:
            dismissal_date = data.order_date
            if data.extra_fields and data.extra_fields.get("dismissal_date"):
                try:
                    from datetime import datetime as _dt
                    dismissal_date = _dt.strptime(data.extra_fields["dismissal_date"], "%Y-%m-%d").date()
                except (ValueError, TypeError):
                    pass
            employee.is_dismissed = True
            employee.dismissal_date = dismissal_date
            employee.dismissal_reason = f"Приказ №{order_number} от {data.order_date.strftime('%d.%m.%Y')}"
            employee.dismissed_by = "system"
            await db.flush()

        audit_logger.info(
            f"ORDER CREATED: number={order_number}, type={order_type.name}, employee_id={data.employee_id}, employee_name={employee.name}",
            extra={
                "employee_id": data.employee_id,
                "employee_name": employee.name,
                "action": "order_created",
                "user_id": "system",
                "order_id": order.id,
                "details": {
                    "order_number": order_number,
                    "order_type_name": order_type.name,
                    "order_type_code": order_type.code,
                    "order_date": str(data.order_date),
                },
            },
        )
        return order

    async def update_order(self, db: AsyncSession, order_id: int, data: OrderUpdate, user_id: str) -> dict[str, Any]:
        order = await self.order_repo.get_by_id(db, order_id)
        if not order:
            raise OrderNotFoundError(f"Приказ {order_id} не найден")

        updates: dict[str, Any] = {}
        if data.order_number is not None:
            updates["order_number"] = data.order_number
        if data.order_date is not None:
            updates["order_date"] = data.order_date
        if data.notes is not None:
            updates["notes"] = data.notes
        if data.extra_fields is not None:
            updates["extra_fields"] = data.extra_fields

        if updates:
            for key, value in updates.items():
                setattr(order, key, value)
            await db.flush()
            await db.refresh(order)

        return self._serialize_order(order)

    async def _generate_document(
        self,
        order_number: str,
        data: OrderCreate,
        employee: Employee,
        order_type: OrderType,
        year_dir: Path,
    ) -> tuple[str, str]:
        doc, replacements = await self._build_document(order_number, data, employee, order_type)

        storage_name = self._build_storage_name(
            order_number, data.order_date, order_type, employee,
            self._extract_extra_dates(data.extra_fields, order_type.code),
        )
        display_name = self._build_display_name(
            order_number, data.order_date, order_type, employee,
            self._extract_extra_info(data.extra_fields, order_type.code),
        )
        file_path = year_dir / storage_name

        if data.draft_id:
            from app.services.order_draft_service import order_draft_service

            draft_path = order_draft_service.get_draft_path(data.draft_id)
            file_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(str(draft_path), str(file_path))
            order_draft_service.delete_draft(data.draft_id)
            return storage_key(file_path, "ORDERS_PATH"), display_name

        await asyncio.wait_for(
            asyncio.to_thread(doc.save, str(file_path)),
            timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
        )
        return storage_key(file_path, "ORDERS_PATH"), display_name

    async def _build_document(
        self,
        order_number: str,
        data: OrderCreate,
        employee: Employee,
        order_type: OrderType,
    ) -> tuple[Document, dict[str, str]]:
        template_path = self._get_template_path(order_type)

        if template_path.exists():
            doc = await asyncio.wait_for(
                asyncio.to_thread(Document, str(template_path)),
                timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
            )
        else:
            doc = Document()
            doc.add_heading(f"Приказ №{order_number}", level=1)
            warning_run = doc.add_paragraph().add_run(MISSING_TEMPLATE_WARNING)
            warning_run.bold = True
            warning_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            doc.add_paragraph(f"Тип: {order_type.name}")
            doc.add_paragraph(f"Дата: {data.order_date.strftime('%d.%m.%Y')}")
            doc.add_paragraph(f"Сотрудник: {employee.name}")

        replacements = self._prepare_replacements(order_number, data, employee, order_type)
        await asyncio.wait_for(
            asyncio.to_thread(self._replace_placeholders, doc, replacements),
            timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
        )
        return doc, replacements

    def _prepare_replacements(
        self,
        order_number: str,
        data: OrderCreate,
        employee: Employee,
        order_type: OrderType,
    ) -> dict[str, str]:
        full_name = employee.name
        name_parts = full_name.split()
        last_name = name_parts[0] if name_parts else "Unknown"
        first_name = name_parts[1] if len(name_parts) > 1 else ""
        middle_name = name_parts[2] if len(name_parts) > 2 else ""
        initials = "_".join([p[0] for p in name_parts[1:]]) if len(name_parts) > 1 else ""
        initials_dots = " ".join([f"{p[0]}." for p in name_parts[1:]]) if len(name_parts) > 1 else ""
        initials_nospace = "".join([f"{p[0]}." for p in name_parts[1:]]) if len(name_parts) > 1 else ""

        position_name = str(employee.position.name if employee.position else "")
        position_cap = position_name.capitalize() if position_name else ""

        # Gender-aware acknowledgment
        oznak = "ознакомлена" if employee.gender == "female" else "ознакомлен"

        replacements = {
            "{order_number}": order_number,
            "{order_date}": data.order_date.strftime("%d.%m.%Y"),
            "{order_type_name}": order_type.name,
            "{order_type_code}": order_type.code,
            "{order_type_lower}": order_type.name.lower(),
            "{full_name}": full_name,
            "{full_name_upper}": full_name.upper(),
            "{full_name_title}": full_name.title(),
            "{full_name_last_caps}": f"{last_name.upper()} {first_name} {middle_name}".strip(),
            "{last_name_upper}": last_name.upper(),
            "{short_name}": f"{last_name} {initials_nospace}".strip(),
            "{initials_before}": f"{initials_nospace}{last_name}".strip(),
            "{last_name_then_initials}": f"{last_name} {initials_nospace}".strip(),
            "{last_name}": last_name,
            "{initials}": initials,
            "{tab_number}": str(employee.tab_number or ""),
            "{department}": str(employee.department.name if employee.department else ""),
            "{position}": position_name.lower(),
            "{position_cap}": position_cap,
            "{hire_date}": employee.hire_date.strftime("%d.%m.%Y") if employee.hire_date else "",
            "{contract_start}": employee.contract_start.strftime("%d.%m.%Y") if employee.contract_start else "",
            "{hire_order_date}": employee.hire_date.strftime("%d.%m.%Y") if employee.hire_date else "",
            "{oznak_gender}": oznak,
            "{notes}": data.notes or "",
        }

        for key, value in (data.extra_fields or {}).items():
            replacements[f"{{{key}}}"] = self._format_placeholder_value(value)

        return replacements

    def _replace_placeholders(self, target: Any, replacements: dict[str, str]) -> None:
        """Replace placeholders in a Document, Paragraph, or Cell. Dispatches to the appropriate handler."""
        if hasattr(target, "paragraphs") and hasattr(target, "tables"):
            # Document
            for paragraph in target.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
            for table in target.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, replacements)
        elif hasattr(target, "paragraphs"):
            # Cell or similar object with paragraphs
            for paragraph in target.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)
        else:
            # Single Paragraph
            self._replace_in_paragraph(target, replacements)

    def _replace_in_paragraph(self, paragraph: Any, replacements: dict[str, str]) -> None:
        """Replace placeholders in a single paragraph using run coordinate mapping.
        All replacements are applied in a single pass (right-to-left) so run indices stay valid."""
        if not paragraph.runs:
            return

        full_text = paragraph.text

        # Collect all occurrences of all keys
        occurrences: list[tuple[int, str, str]] = []  # (start_pos, key, value)
        for key, value in replacements.items():
            if key not in full_text:
                continue
            key_len = len(key)
            for i in range(len(full_text)):
                if full_text.startswith(key, i):
                    occurrences.append((i, key, value))

        if not occurrences:
            return

        # Sort by position descending (right-to-left) so earlier indices stay valid
        occurrences.sort(key=lambda x: x[0], reverse=True)

        # Build coordinate map once
        p_map = []
        for run_idx, run in enumerate(paragraph.runs):
            for char_idx, char in enumerate(run.text):
                p_map.append({"run": run_idx, "char": char_idx})

        for start_pos, key, value in occurrences:
            key_len = len(key)
            key_map = p_map[start_pos : start_pos + key_len]
            self._replace_in_runs(paragraph.runs, key_map, value)

    def _replace_in_runs(self, runs: list[Any], key_map: list[dict], value: str) -> None:
        """Apply replacement to the specific runs/characters identified by key_map."""
        for i, position in enumerate(reversed(key_map), start=1):
            run_idx = position["run"]
            char_idx = position["char"]
            run = runs[run_idx]
            chars = list(run.text)

            if i < len(key_map):
                # Not the first character of the key — delete it
                chars.pop(char_idx)
            else:
                # First character (last in reversed order) — replace with value
                chars[char_idx] = value

            run.text = "".join(chars)

    async def sync_orders(self, db: AsyncSession, year: Optional[int] = None) -> dict[str, Any]:
        years_to_check = [year] if year else await self.order_repo.get_years(db)
        if not years_to_check:
            return {"message": "Нет данных для синхронизации", "deleted": 0, "added": 0}

        deleted = 0
        for y in years_to_check:
            year_dir = Path(settings.ORDERS_PATH) / str(y)
            if not year_dir.exists():
                continue

            files_on_disk = {f.name for f in year_dir.iterdir() if f.is_file() and f.suffix == ".docx"}
            orders_in_db = await self.order_repo.get_all(db, page=1, per_page=10000, year=y)
            db_files = {Path(storage_key(o.file_path, "ORDERS_PATH")).name for o in orders_in_db[0] if o.file_path}
            missing_files = db_files - files_on_disk
            for order in orders_in_db[0]:
                if order.file_path and Path(storage_key(order.file_path, "ORDERS_PATH")).name in missing_files:
                    await self.order_repo.soft_delete(db, order.id, "sync")
                    deleted += 1

        await db.commit()
        return {"message": f"Синхронизация завершена: удалено {deleted}, добавление новых файлов отключено", "deleted": deleted, "added": 0}

    async def upload_template(self, db: AsyncSession, order_type_id: int, filename: str, content: bytes) -> dict[str, Any]:
        order_type = await self.get_order_type(db, order_type_id)
        storage_name = self._normalize_template_filename(filename, order_type.code)
        display_name = f"Шаблон - {order_type.name}.docx"
        template_path = Path(settings.TEMPLATES_PATH) / storage_name
        template_path.parent.mkdir(parents=True, exist_ok=True)
        with open(template_path, "wb") as file_obj:
            file_obj.write(content)
        await self.order_type_repo.update(db, order_type, {
            "template_filename": storage_name,
            "display_name": display_name,
        })
        await db.commit()
        return self._serialize_order_type(order_type)

    async def delete_template(self, db: AsyncSession, order_type_id: int) -> None:
        order_type = await self.get_order_type(db, order_type_id)
        template_path = self._get_template_path(order_type)
        if template_path.exists():
            template_path.unlink()
        await self.order_type_repo.update(db, order_type, {"template_filename": None})
        await db.commit()

    async def bulk_upload_templates(self, db: AsyncSession, files: list[Any]) -> dict[str, Any]:
        results: dict[str, Any] = {"uploaded": 0, "skipped": 0, "errors": []}
        for file in files:
            if not file.filename or not file.filename.endswith(".docx"):
                results["skipped"] += 1
                results["errors"].append(f"{file.filename or 'unknown'}: только .docx")
                continue
            code = Path(file.filename).stem
            order_type = await self.order_type_repo.get_by_code(db, code)
            if not order_type:
                results["skipped"] += 1
                results["errors"].append(f"{file.filename}: тип приказа с кодом '{code}' не найден")
                continue
            try:
                content = await file.read()
                safe_filename = self._normalize_template_filename(file.filename, order_type.code)
                template_path = Path(settings.TEMPLATES_PATH) / safe_filename
                template_path.parent.mkdir(parents=True, exist_ok=True)
                with open(template_path, "wb") as file_obj:
                    file_obj.write(content)
                await self.order_type_repo.update(db, order_type, {"template_filename": safe_filename})
                results["uploaded"] += 1
            except Exception as e:
                results["errors"].append(f"{file.filename}: {str(e)}")
                results["skipped"] += 1
        await db.commit()
        return results

    def get_template_variables(self) -> list[dict[str, str]]:
        """Возвращает список всех доступных переменных для шаблонов с описаниями"""
        return [
            {"name": "{order_number}", "description": "Номер приказа", "category": "Приказ"},
            {"name": "{order_date}", "description": "Дата приказа (ДД.ММ.ГГГГ)", "category": "Приказ"},
            {"name": "{order_type_name}", "description": "Название типа приказа", "category": "Приказ"},
            {"name": "{order_type_code}", "description": "Код типа приказа", "category": "Приказ"},
            {"name": "{order_type_lower}", "description": "Тип приказа строчными буквами", "category": "Приказ"},

            {"name": "{full_name}", "description": "ФИО полностью", "category": "ФИО"},
            {"name": "{full_name_upper}", "description": "ФИО заглавными буквами", "category": "ФИО"},
            {"name": "{full_name_title}", "description": "ФИО с заглавной буквы", "category": "ФИО"},
            {"name": "{full_name_last_caps}", "description": "Фамилия заглавными, имя отчество обычными", "category": "ФИО"},
            {"name": "{last_name_upper}", "description": "Фамилия заглавными буквами", "category": "ФИО"},
            {"name": "{short_name}", "description": "Фамилия И.О.", "category": "ФИО"},
            {"name": "{initials_before}", "description": "И.О.Фамилия (без пробелов)", "category": "ФИО"},
            {"name": "{last_name_then_initials}", "description": "Фамилия И.О. (без пробела)", "category": "ФИО"},
            {"name": "{last_name}", "description": "Фамилия", "category": "ФИО"},
            {"name": "{initials}", "description": "Инициалы через подчеркивание (для имени файла)", "category": "ФИО"},

            {"name": "{position}", "description": "Должность (все строчные)", "category": "Работа"},
            {"name": "{position_cap}", "description": "Должность (с заглавной буквы)", "category": "Работа"},
            {"name": "{department}", "description": "Подразделение", "category": "Работа"},
            {"name": "{tab_number}", "description": "Табельный номер", "category": "Работа"},

            {"name": "{hire_date}", "description": "Дата приема на работу (из карточки сотрудника)", "category": "Даты"},
            {"name": "{contract_start}", "description": "Дата начала контракта", "category": "Даты"},
            {"name": "{contract_end}", "description": "Дата окончания контракта (вводится вручную)", "category": "Даты"},
            {"name": "{trial_end}", "description": "Дата окончания испытательного срока (вводится вручную)", "category": "Даты"},
            {"name": "{hire_order_date}", "description": "Дата приема (для приказа «Прием на работу»)", "category": "Даты"},
            {"name": "{dismissal_date}", "description": "Дата увольнения (для приказа «Увольнение»)", "category": "Даты"},
            {"name": "{vacation_start}", "description": "Начало отпуска (для приказов «Отпуск»)", "category": "Даты"},
            {"name": "{vacation_end}", "description": "Конец отпуска (для приказов «Отпуск»)", "category": "Даты"},
            {"name": "{vacation_days}", "description": "Кол-во дней отпуска", "category": "Даты"},
            {"name": "{sick_leave_start}", "description": "Начало больничного", "category": "Даты"},
            {"name": "{sick_leave_end}", "description": "Конец больничного", "category": "Даты"},
            {"name": "{sick_leave_days}", "description": "Кол-во дней больничного", "category": "Даты"},
            {"name": "{transfer_date}", "description": "Дата перевода", "category": "Даты"},
            {"name": "{contract_new_end}", "description": "Новая дата конца контракта (для «Продление контракта»)", "category": "Даты"},
            {"name": "{call_date}", "description": "Дата вызова (для «Вызов в выходной»)", "category": "Даты"},
            {"name": "{call_date_start}", "description": "Дата начала вызова (для «Вызов в выходной»)", "category": "Даты"},
            {"name": "{call_date_end}", "description": "Дата окончания вызова (для «Вызов в выходной»)", "category": "Даты"},
            {"name": "{recall_date}", "description": "Дата отзыва из отпуска (для «Отзыв из отпуска»)", "category": "Даты"},

            {"name": "{oznak_gender}", "description": "ознакомлен/ознакомлена (по полу сотрудника, с маленькой буквы)", "category": "Прочее"},
            {"name": "{notes}", "description": "Комментарий к приказу", "category": "Прочее"},

            {"name": "{<key из field_schema>}", "description": "Любое дополнительное поле типа приказа", "category": "Поля типа"},
        ]

    async def cancel_order(self, db: AsyncSession, order_id: int, user_id: str) -> bool:
        order = await self.order_repo.get_by_id(db, order_id)
        if not order:
            raise OrderNotFoundError(order_id)
        await self.order_repo.cancel(db, order_id, user_id)

        # Восстановление сотрудника при отмене приказа об увольнении
        order_type = order.order_type
        if not order_type:
            from sqlalchemy import select as _select
            from app.models.order_type import OrderType as _OrderType
            type_result = await db.execute(
                _select(_OrderType).where(_OrderType.id == order.order_type_id)
            )
            order_type = type_result.scalar_one_or_none()
        if order_type and order_type.code == "dismissal":
            employee = await self.employee_repo.get_by_id(db, order.employee_id)
            if employee and employee.is_dismissed:
                employee.is_dismissed = False
                employee.dismissal_date = None
                employee.dismissal_reason = None
                employee.dismissed_by = None
                employee.dismissed_at = None
                await db.flush()

        # Восстановление contract_end при отмене приказа о продлении
        if order_type and order_type.code == "contract_extension":
            employee = await self.employee_repo.get_by_id(db, order.employee_id)
            if employee and order.extra_fields and order.extra_fields.get("old_contract_end"):
                employee.contract_end = date.fromisoformat(order.extra_fields["old_contract_end"])
                await db.flush()

        await db.commit()
        return True

    async def hard_delete_order(self, db: AsyncSession, order_id: int) -> bool:
        order = await self.order_repo.get_by_id(db, order_id)
        if not order:
            raise OrderNotFoundError(order_id)

        employee = await self.employee_repo.get_by_id(db, order.employee_id)

        # Восстановление сотрудника при удалении приказа об увольнении
        order_type_code = order.order_type.code if getattr(order, "order_type", None) else None
        if not order_type_code:
            from sqlalchemy import select as _select
            from app.models.order_type import OrderType as _OrderType
            type_result = await db.execute(
                _select(_OrderType.code).where(_OrderType.id == order.order_type_id)
            )
            order_type_code = type_result.scalar_one_or_none()
        if order_type_code == "dismissal" and employee and employee.is_dismissed:
            employee.is_dismissed = False
            employee.dismissal_date = None
            employee.dismissal_reason = None
            employee.dismissed_by = None
            employee.dismissed_at = None
            await db.flush()

        # Восстановление contract_end при удалении приказа о продлении
        if order_type_code == "contract_extension" and employee and order.extra_fields and order.extra_fields.get("old_contract_end"):
            employee.contract_end = date.fromisoformat(order.extra_fields["old_contract_end"])
            await db.flush()

        # Удаляем приказ точечно: запоминаем затронутые period_id и после CASCADE
        # пересчитываем только их по оставшемуся журналу операций.
        # Важно: не запускаем recalculate_periods, чтобы не перераспределять
        # все отпуска "с нуля" по другим периодам.
        from app.services.vacation_period_service import vacation_period_service

        affected_period_ids = await vacation_period_service.get_affected_period_ids_for_order(db, order_id)

        # Чистим зависимые сущности явно, чтобы удаление было устойчивым даже если
        # каскадные FK ещё не применены в конкретной БД.
        from app.models.vacation import Vacation
        from app.models.vacation_adjustment import VacationAdjustment
        from app.models.vacation_period_manual_closure import VacationPeriodManualClosure
        from app.models.vacation_period_transaction import VacationPeriodTransaction

        tx_ids_result = await db.execute(
            select(VacationPeriodTransaction.id).where(
                or_(
                    VacationPeriodTransaction.original_order_id == order_id,
                    VacationPeriodTransaction.adjustment_order_id == order_id,
                )
            )
        )
        tx_ids = [row[0] for row in tx_ids_result.all()]
        if tx_ids:
            await db.execute(
                update(VacationPeriodTransaction)
                .where(VacationPeriodTransaction.reversed_transaction_id.in_(tx_ids))
                .values(reversed_transaction_id=None)
            )
            await db.execute(
                sa_delete(VacationPeriodTransaction).where(VacationPeriodTransaction.id.in_(tx_ids))
            )

        vacation_ids_result = await db.execute(
            select(Vacation.id).where(
                or_(
                    Vacation.order_id == order_id,
                    Vacation.recall_order_id == order_id,
                    Vacation.postpone_order_id == order_id,
                    Vacation.extension_order_id == order_id,
                )
            )
        )
        vacation_ids = [row[0] for row in vacation_ids_result.all()]
        if vacation_ids:
            await db.execute(
                update(VacationPeriodTransaction)
                .where(VacationPeriodTransaction.vacation_id.in_(vacation_ids))
                .values(vacation_id=None)
            )
            await db.execute(sa_delete(Vacation).where(Vacation.id.in_(vacation_ids)))

        await db.execute(
            sa_delete(VacationAdjustment).where(
                or_(
                    VacationAdjustment.original_order_id == order_id,
                    VacationAdjustment.adjustment_order_id == order_id,
                )
            )
        )
        await db.execute(sa_delete(VacationPeriodManualClosure).where(VacationPeriodManualClosure.order_id == order_id))

        # Удаляем файл приказа
        if order.file_path:
            try:
                storage_path(order.file_path, "ORDERS_PATH").unlink()
            except OSError:
                pass

        # Удаляем приказ — CASCADE удалит связанные отпуска, транзакции и корректировки
        await self.order_repo.hard_delete(db, order_id)
        await vacation_period_service.recompute_period_totals_by_ids(db, affected_period_ids)
        await db.commit()

        audit_logger.info(
            f"ORDER DELETED: id={order_id}, number={order.order_number}, type={order.order_type.name if order.order_type else ''}, employee_id={order.employee_id}, employee_name={employee.name if employee else None}",
            extra={"employee_id": order.employee_id, "employee_name": employee.name if employee else None, "action": "order_deleted", "user_id": "system", "order_id": order_id},
        )
        return True

    async def create_vacation_unpaid_group_order(
        self, db: AsyncSession, data: "VacationUnpaidGroupOrderCreate"
    ) -> Order:
        """Create a group unpaid vacation order for multiple employees."""
        from datetime import timedelta
        from sqlalchemy import insert as sa_insert

        from app.models.order_employee import OrderEmployee
        from app.models.vacation import Vacation
        from app.repositories.vacation_repository import vacation_repository as _vacation_repo

        await self.ensure_default_order_types(db)

        if not data.employees:
            raise HRMSException("Список сотрудников не может быть пустым", "validation_error", status_code=422)

        order_type = await self.get_order_type_by_code(db, "vacation_unpaid")

        order_number = data.order_number
        if not order_number:
            order_number = await self.order_repo.get_next_order_number(db, order_type.id)
        else:
            order_number = order_number.strip()

        common_start = data.vacation_start

        employee_rows = []
        for emp_item in data.employees:
            employee = await self.employee_repo.get_by_id(db, emp_item.employee_id)
            if not employee:
                raise EmployeeNotFoundError(emp_item.employee_id)

            vacation_end = common_start + timedelta(days=emp_item.vacation_days - 1)

            overlap = await _vacation_repo.check_overlap(db, emp_item.employee_id, common_start, vacation_end)
            if overlap:
                raise VacationOverlapError(
                    f"Пересечение отпуска для сотрудника {employee.name}: {overlap.start_date} — {overlap.end_date}"
                )

            employee_rows.append({
                "employee": employee,
                "vacation_days": emp_item.vacation_days,
                "vacation_end": vacation_end,
            })

        year_dir = Path(settings.ORDERS_PATH) / str(data.order_date.year)
        year_dir.mkdir(parents=True, exist_ok=True)

        file_path, display_name = await self._generate_group_document(
            order_number, data, order_type, year_dir, employee_rows,
        )

        order = await self.order_repo.create(
            db,
            {
                "order_number": order_number,
                "order_type_id": order_type.id,
                "employee_id": None,
                "order_date": data.order_date,
                "file_path": file_path,
                "display_name": display_name,
                "notes": None,
                "extra_fields": {},
                "is_group": True,
            },
        )

        for row in employee_rows:
            emp = row["employee"]
            await db.execute(
                sa_insert(OrderEmployee).values(
                    order_id=order.id,
                    employee_id=emp.id,
                    vacation_start=common_start,
                    vacation_end=row["vacation_end"],
                    vacation_days=row["vacation_days"],
                )
            )
            await db.execute(
                sa_insert(Vacation).values(
                    employee_id=emp.id,
                    start_date=common_start,
                    end_date=row["vacation_end"],
                    vacation_type="Отпуск за свой счет",
                    days_count=row["vacation_days"],
                    vacation_year=common_start.year,
                    order_id=order.id,
                )
            )

        await db.flush()

        # Reload order with relationships for serialization
        from sqlalchemy import select as sa_select
        from sqlalchemy.orm import selectinload
        from app.models.order_employee import OrderEmployee
        reload_result = await db.execute(
            sa_select(Order)
            .options(
                selectinload(Order.order_type),
                selectinload(Order.employees).selectinload(OrderEmployee.employee),
            )
            .where(Order.id == order.id)
            .execution_options(populate_existing=True)
        )
        order = reload_result.scalar_one()

        audit_logger.info(
            f"GROUP ORDER CREATED: number={order_number}, type=vacation_unpaid, employee_count={len(employee_rows)}",
            extra={
                "action": "group_order_created",
                "user_id": "system",
                "order_id": order.id,
                "details": {
                    "order_number": order_number,
                    "order_type_code": "vacation_unpaid",
                    "order_date": str(data.order_date),
                    "employee_count": len(employee_rows),
                },
            },
        )

        return order

    async def _generate_group_document(
        self,
        order_number: str,
        data: "VacationUnpaidGroupOrderCreate",
        order_type: OrderType,
        year_dir: Path,
        employee_rows: list[dict],
    ) -> tuple[str, str]:
        """Generate DOCX for a group order using the group template."""
        template_name = "template__order__vacation_unpaid_group.docx"
        template_path = Path(settings.TEMPLATES_PATH) / template_name

        if template_path.exists():
            doc = await asyncio.wait_for(
                asyncio.to_thread(Document, str(template_path)),
                timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
            )
        else:
            doc = Document()
            doc.add_heading(f"Приказ №{order_number}", level=1)
            warning_run = doc.add_paragraph().add_run(MISSING_TEMPLATE_WARNING)
            warning_run.bold = True
            warning_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        replacements = {
            "{order_number}": order_number,
            "{order_date}": data.order_date.strftime("%d.%m.%Y"),
            "{vacation_start}": data.vacation_start.strftime("%d.%m.%Y"),
        }

        employee_table_replacements = []
        for idx, row_data in enumerate(employee_rows, 1):
            emp = row_data["employee"]
            employee_table_replacements.append({
                "{index}": str(idx),
                "{full_name}": emp.name,
                "{position}": str(emp.position.name if emp.position else ""),
                "{department}": str(emp.department.name if emp.department else ""),
                "{vacation_start}": data.vacation_start.strftime("%d.%m.%Y"),
                "{vacation_end}": row_data["vacation_end"].strftime("%d.%m.%Y"),
                "{vacation_days}": str(row_data["vacation_days"]),
            })

        # Process employee table rows BEFORE general placeholder replacement
        for table in doc.tables:
            template_row = None
            for row in table.rows:
                cell_text = " ".join(cell.text for cell in row.cells)
                if "{full_name}" in cell_text:
                    template_row = row
                    break

            if template_row:
                template_idx = None
                for i, row in enumerate(table.rows):
                    if row is template_row:
                        template_idx = i
                        break
                if template_idx is None:
                    continue

                # First add all extra rows from template (before replacing placeholders)
                for _ in employee_table_replacements[1:]:
                    table.add_row()

                # Now replace placeholders in each row starting from template_idx
                for emp_idx, emp_data in enumerate(employee_table_replacements):
                    target_row = table.rows[template_idx + emp_idx]
                    for cell in target_row.cells:
                        for placeholder, value in emp_data.items():
                            self._replace_placeholders(cell.paragraphs[0], {placeholder: value})

        self._replace_placeholders(doc, replacements)

        storage_name = f"prikaz_{order_number}_vacation_unpaid_group_{data.vacation_start.strftime('%Y-%m-%d')}.docx"
        display_name = f"Приказ №{order_number} от {data.order_date.strftime('%d.%m.%Y')} — отпуск за свой счет (групповой, {len(employee_rows)} сотр.)"
        file_path = year_dir / storage_name

        await asyncio.wait_for(
            asyncio.to_thread(doc.save, str(file_path)),
            timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
        )
        return storage_key(file_path, "ORDERS_PATH"), display_name


    def _serialize_order(self, order: Order) -> dict[str, Any]:
        from sqlalchemy import inspect as sa_inspect
        from sqlalchemy.orm.attributes import LoaderCallableStatus

        state = sa_inspect(order)
        loaded = state.attrs

        def _get_loaded(name: str):
            attr = getattr(loaded, name, None)
            if attr is None:
                return None
            val = attr.loaded_value
            if val is LoaderCallableStatus:
                return None
            return val

        ot = _get_loaded("order_type")
        emp = _get_loaded("employee")
        emps = _get_loaded("employees")

        order_type_name = ot.name if ot else ""
        order_type_code = ot.code if ot else ""
        employee_name = emp.name if emp else None
        is_group = bool(getattr(order, "is_group", False))
        group_employee_count = len(emps) if is_group and emps else None

        group_employees = None
        if is_group and emps:
            group_employees = [
                {
                    "employee_id": e.employee_id,
                    "employee_full_name": e.employee.name if e.employee else None,
                    "position": e.employee.position.name if e.employee and e.employee.position else None,
                    "department": e.employee.department.name if e.employee and e.employee.department else None,
                    "vacation_start": e.vacation_start.isoformat() if e.vacation_start else None,
                    "vacation_end": e.vacation_end.isoformat() if e.vacation_end else None,
                    "vacation_days": e.vacation_days,
                }
                for e in emps
            ]
        
        return {
            "id": order.id,
            "order_number": order.order_number,
            "order_type_id": order.order_type_id,
            "order_type_name": order_type_name,
            "order_type_code": order_type_code,
            "employee_id": order.employee_id,
            "employee_name": employee_name,
            "order_date": order.order_date,
            "created_date": order.created_date,
            "file_path": order.file_path,
            "display_name": order.display_name,
            "notes": order.notes,
            "extra_fields": order.extra_fields or {},
            "is_group": is_group,
            "group_employee_count": group_employee_count,
            "group_employees": group_employees,
        }

    def _serialize_order_type(self, order_type: OrderType) -> dict[str, Any]:
        template_path = self._get_template_path(order_type)
        result = {
            "id": order_type.id,
            "code": order_type.code,
            "name": order_type.name,
            "is_active": order_type.is_active,
            "show_in_orders_page": order_type.show_in_orders_page,
            "template_filename": order_type.template_filename,
            "display_name": order_type.display_name,
            "field_schema": order_type.field_schema or [],
            "filename_pattern": order_type.filename_pattern,
            "letter": order_type.letter,
            "template_exists": template_path.exists(),
            "created_at": order_type.created_at,
            "updated_at": order_type.updated_at,
        }
        if template_path.exists():
            stat = template_path.stat()
            result["file_size"] = stat.st_size
            result["last_modified"] = datetime.fromtimestamp(stat.st_mtime).isoformat()
        else:
            result["file_size"] = None
            result["last_modified"] = None
        return result

    def _get_template_path(self, order_type: OrderType) -> Path:
        if not order_type.template_filename:
            return Path(settings.TEMPLATES_PATH) / "__missing__.docx"
        return storage_path(order_type.template_filename, "TEMPLATES_PATH")

    def _extract_extra_dates(self, extra_fields: dict | None, order_type_code: str) -> list[date] | None:
        """Extract relevant dates from extra_fields for storage filename."""
        if not extra_fields:
            return None
        dates: list[date] = []
        try:
            if order_type_code in ("vacation_paid", "vacation_unpaid"):
                if extra_fields.get("vacation_start"):
                    dates.append(date.fromisoformat(extra_fields["vacation_start"]))
                if extra_fields.get("vacation_end"):
                    dates.append(date.fromisoformat(extra_fields["vacation_end"]))
            elif order_type_code == "vacation_recall":
                if extra_fields.get("recall_date"):
                    dates.append(date.fromisoformat(extra_fields["recall_date"]))
            elif order_type_code == "vacation_postpone":
                if extra_fields.get("new_vacation_start"):
                    dates.append(date.fromisoformat(extra_fields["new_vacation_start"]))
                if extra_fields.get("new_vacation_end"):
                    dates.append(date.fromisoformat(extra_fields["new_vacation_end"]))
            elif order_type_code == "vacation_extension":
                if extra_fields.get("sick_start_date"):
                    dates.append(date.fromisoformat(extra_fields["sick_start_date"]))
                if extra_fields.get("sick_end_date"):
                    dates.append(date.fromisoformat(extra_fields["sick_end_date"]))
        except (ValueError, TypeError):
            pass
        return dates or None

    def _extract_extra_info(self, extra_fields: dict | None, order_type_code: str) -> str:
        """Extract human-readable extra info for display name."""
        if not extra_fields:
            return ""
        parts: list[str] = []
        if order_type_code in ("vacation_paid", "vacation_unpaid"):
            start = extra_fields.get("vacation_start", "")
            end = extra_fields.get("vacation_end", "")
            if start and end:
                try:
                    s = date.fromisoformat(start).strftime("%d.%m.%Y")
                    e = date.fromisoformat(end).strftime("%d.%m.%Y")
                    parts.append(f"{s}-{e}")
                except ValueError:
                    pass
        elif order_type_code == "vacation_recall":
            recall = extra_fields.get("recall_date", "")
            if recall:
                try:
                    parts.append(f"с {date.fromisoformat(recall).strftime('%d.%m.%Y')}")
                except ValueError:
                    pass
        elif order_type_code == "vacation_postpone":
            ns = extra_fields.get("new_vacation_start", "")
            ne = extra_fields.get("new_vacation_end", "")
            if ns and ne:
                try:
                    parts.append(f"{date.fromisoformat(ns).strftime('%d.%m.%Y')}-{date.fromisoformat(ne).strftime('%d.%m.%Y')}")
                except ValueError:
                    pass
        elif order_type_code == "vacation_extension":
            ss = extra_fields.get("sick_start_date", "")
            se = extra_fields.get("sick_end_date", "")
            if ss and se:
                try:
                    parts.append(f"больничный {date.fromisoformat(ss).strftime('%d.%m.%Y')}-{date.fromisoformat(se).strftime('%d.%m.%Y')}")
                except ValueError:
                    pass
        return " ".join(parts)

    @staticmethod
    def _transliterate(text: str) -> str:
        """Convert Cyrillic text to Latin transcription."""
        mapping = {
            'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'yo',
            'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
            'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
            'ф': 'f', 'х': 'kh', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'shch',
            'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
            'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'Yo',
            'Ж': 'Zh', 'З': 'Z', 'И': 'I', 'Й': 'Y', 'К': 'K', 'Л': 'L', 'М': 'M',
            'Н': 'N', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U',
            'Ф': 'F', 'Х': 'Kh', 'Ц': 'Ts', 'Ч': 'Ch', 'Ш': 'Sh', 'Щ': 'Shch',
            'Ъ': '', 'Ы': 'Y', 'Ь': '', 'Э': 'E', 'Ю': 'Yu', 'Я': 'Ya',
        }
        result = []
        for ch in text:
            result.append(mapping.get(ch, ch))
        return ''.join(result)

    def _build_storage_name(
        self,
        order_number: str,
        order_date: date,
        order_type: OrderType,
        employee: Employee,
        extra_dates: list[date] | None = None,
    ) -> str:
        """Build a filesystem-safe filename (ASCII only, no spaces)."""
        type_code = order_type.code
        last_name = employee.name.split()[0] if employee.name else "unknown"
        transliterated = self._transliterate(last_name).lower()
        transliterated = re.sub(r'[^a-z0-9-]', '', transliterated.replace(' ', '-'))

        parts = [
            order_date.isoformat(),
            f"prikaz_{order_number}",
            type_code,
            transliterated,
        ]
        if extra_dates:
            date_strs = [d.isoformat() for d in extra_dates]
            parts.append("_".join(date_strs))
        return "_".join(parts) + ".docx"

    def _build_display_name(
        self,
        order_number: str,
        order_date: date,
        order_type: OrderType,
        employee: Employee,
        extra_info: str = "",
    ) -> str:
        """Build a human-readable display name in Russian."""
        date_str = order_date.strftime("%d.%m.%Y")
        name = f"Приказ №{order_number} от {date_str} - {order_type.name} - {employee.name}"
        if extra_info:
            name += f" - {extra_info}"
        return name + ".docx"

    def _normalize_template_filename(self, original_filename: str, code: str) -> str:
        ext = Path(original_filename).suffix or ".docx"
        return f"template__order__{code}{ext.lower()}"

    def _build_filename(self, order_number: str, order_type: OrderType, replacements: dict[str, str]) -> str:
        pattern = order_type.filename_pattern or "Приказ_№{order_number}_{order_type_code}_{last_name}_{initials}.docx"
        filename = pattern
        for key, value in replacements.items():
            filename = filename.replace(key, value)
        if not filename.lower().endswith(".docx"):
            filename = f"{filename}.docx"
        sanitized = re.sub(r'[<>:"/\\\\|?*]+', "_", filename).strip()
        return sanitized or f"order_{order_number}.docx"

    def _format_placeholder_value(self, value: Any) -> str:
        if isinstance(value, str):
            try:
                parsed = datetime.strptime(value, "%Y-%m-%d")
                return parsed.strftime("%d.%m.%Y")
            except ValueError:
                return value
        return str(value)


order_service = OrderService()
