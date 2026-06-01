"""Extract placeholders from DOCX templates and suggest field schemas."""
import re
from pathlib import Path
from typing import Any

from docx import Document

from .template_variables_service import ALL_TEMPLATE_VARIABLES

# Плейсхолдеры которые НЕ нужно показывать в форме — они заполняются автоматически
AUTO_FILLED_PLACEHOLDERS = frozenset({
    # Employee data
    "full_name", "short_name", "last_name", "first_name", "middle_name",
    "full_name_upper", "full_name_title", "full_name_last_caps",
    "last_name_upper", "initials_before", "last_name_then_initials", "initials",
    "position", "position_cap", "department", "tab_number",
    "hire_date", "contract_start", "oznak", "oznak_gender",
    # Document base
    "doc_number", "doc_date", "doc_title",
    # Order specific
    "order_number", "order_date", "order_type_name", "order_type_code", "order_type_lower",
    "hire_order_date",
    # Notification specific
    "notification_type_name", "notification_type_code",
    # Statement specific
    "statement_type_name", "statement_type_code",
    # Calculated
    "trial_end_months", "contract_end_years", "new_contract_years",
    # Contract extension (auto-filled from employee)
    "old_contract_start", "old_contract_end",
    # Block markers
    "employees_block_start", "employees_block_end",
    "applications_block_start", "applications_block_end",
    # Other
    "index", "notes",
})

# Паттерн для поиска плейсхолдеров
PLACEHOLDER_RE = re.compile(r"\{([a-zA-Z_][a-zA-Z0-9_]*)\}")


def extract_placeholders_from_docx(file_path: Path) -> list[str]:
    """Extract all unique placeholder names from a DOCX file."""
    if not file_path.exists():
        return []

    doc = Document(str(file_path))
    placeholders: set[str] = set()

    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        placeholders.update(PLACEHOLDER_RE.findall(paragraph.text))

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.update(PLACEHOLDER_RE.findall(cell.text))

    return sorted(placeholders)


def _key_to_display_name(key: str) -> str | None:
    """Get short displayName for a placeholder key."""
    for tv in ALL_TEMPLATE_VARIABLES:
        if tv["name"].strip("{}") == key:
            return tv.get("displayName")
    return None


def _key_to_label(key: str) -> str:
    """Convert snake_case key to human-readable Russian label."""
    for tv in ALL_TEMPLATE_VARIABLES:
        if tv["name"].strip("{}") == key:
            return tv.get("displayName", tv["description"])
    # Fallback: snake_case → Title Case
    return key.replace("_", " ").title()


def _suggest_field_type(key: str) -> str:
    """Suggest field type based on key name."""
    key_lower = key.lower()

    # Date fields — match suffix/standalone, not substring
    date_words = ("date", "start", "end", "recall", "trial")
    if key_lower.endswith(date_words) or key_lower in date_words:
        return "date"

    # Number fields
    if any(word in key_lower for word in ("days", "months", "years", "count", "number")):
        return "number"

    # Textarea for long content
    if any(word in key_lower for word in ("comment", "note", "reason", "description")):
        return "textarea"

    return "text"


def suggest_field_schema(placeholders: list[str]) -> list[dict[str, Any]]:
    """Build field schema suggestions from extracted placeholders.

    ALL placeholders from template are included so user can see and edit them.
    Fields are intelligently arranged into a 2-column grid layout:
    - start/end date pairs are placed on the same row
    - number fields (days, years) are placed next to their related dates
    - remaining fields are arranged sequentially
    """
    WIDTH_BY_TYPE = {"text": 250, "number": 100, "date": 130, "textarea": None}

    # Categorize fields — two-pass approach to avoid duplicates
    starts = {p for p in placeholders if p.endswith("_start")}
    ends = {p for p in placeholders if p.endswith("_end")}

    date_pairs: list[tuple[str, str]] = []
    paired = set()

    # First pass: find all start/end pairs
    for start in starts:
        counterpart = start[:-6] + "_end"
        if counterpart in ends:
            date_pairs.append((start, counterpart))
            paired.add(start)
            paired.add(counterpart)

    # Second pass: everything else
    standalone_dates: list[str] = []
    numbers: list[str] = []
    texts: list[str] = []

    for placeholder in placeholders:
        if placeholder in paired:
            continue
        field_type = _suggest_field_type(placeholder)
        if field_type == "date":
            standalone_dates.append(placeholder)
        elif field_type == "number":
            numbers.append(placeholder)
        else:
            texts.append(placeholder)

    # Fields that auto-fill from employee/order data — hide by default
    AUTO_HIDDEN = frozenset({
        # Employee data
        "full_name", "short_name", "last_name", "first_name", "middle_name",
        "full_name_upper", "full_name_title", "full_name_last_caps",
        "last_name_upper", "initials_before", "last_name_then_initials", "initials",
        "position", "position_cap", "department", "tab_number",
        # Order base data
        "order_number", "order_date", "order_type_name", "order_type_code", "order_type_lower",
        # Other auto-filled
        "oznak", "oznak_gender",
    })

    # Build schema with smart positioning
    schema = []
    row = 0
    used = set()

    # Place date pairs on the same row
    for start, end in date_pairs:
        schema.append({
            "key": start,
            "label": _key_to_label(start),
            "displayName": _key_to_display_name(start),
            "type": "date",
            "required": False,
            "enabled": start not in AUTO_HIDDEN,
            "col": 0,
            "row": row,
            "width": WIDTH_BY_TYPE["date"],
        })
        schema.append({
            "key": end,
            "label": _key_to_label(end),
            "displayName": _key_to_display_name(end),
            "type": "date",
            "required": False,
            "enabled": end not in AUTO_HIDDEN,
            "col": 1,
            "row": row,
            "width": WIDTH_BY_TYPE["date"],
        })
        used.add(start)
        used.add(end)
        row += 1

    # Place standalone dates, try to pair with related numbers
    for date_key in standalone_dates:
        base = date_key.rsplit("_", 1)[0] if "_" in date_key else date_key
        related_num = next((n for n in numbers if n not in used and (base in n or n.startswith(base))), None)

        schema.append({
            "key": date_key,
            "label": _key_to_label(date_key),
            "displayName": _key_to_display_name(date_key),
            "type": "date",
            "required": False,
            "enabled": date_key not in AUTO_HIDDEN,
            "col": 0,
            "row": row,
            "width": WIDTH_BY_TYPE["date"],
        })
        used.add(date_key)

        if related_num:
            schema.append({
                "key": related_num,
                "label": _key_to_label(related_num),
                "displayName": _key_to_display_name(related_num),
                "type": "number",
                "required": False,
                "enabled": related_num not in AUTO_HIDDEN,
                "col": 1,
                "row": row,
                "width": WIDTH_BY_TYPE["number"],
            })
            used.add(related_num)

        row += 1

    # Place remaining numbers
    for num_key in numbers:
        if num_key not in used:
            schema.append({
                "key": num_key,
                "label": _key_to_label(num_key),
                "displayName": _key_to_display_name(num_key),
                "type": "number",
                "required": False,
                "enabled": num_key not in AUTO_HIDDEN,
                "col": 0,
                "row": row,
                "width": WIDTH_BY_TYPE["number"],
            })
            used.add(num_key)
            row += 1

    # Place text fields in 2 columns
    text_row = row
    for idx, text_key in enumerate(texts):
        col = idx % 2
        if col == 0 and idx > 0:
            text_row += 1
        schema.append({
            "key": text_key,
            "label": _key_to_label(text_key),
            "displayName": _key_to_display_name(text_key),
            "type": "text",
            "required": False,
            "enabled": text_key not in AUTO_HIDDEN,
            "col": col,
            "row": text_row,
            "width": WIDTH_BY_TYPE["text"],
        })
        used.add(text_key)

    return schema
