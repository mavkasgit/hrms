"""Shared DOCX placeholder renderer for all document types.

Provides a generic function to replace {placeholder} tokens in DOCX files,
reusing the same run-coordinate-aware algorithm as the order document service.
"""
import asyncio
from pathlib import Path
from typing import Any

from docx import Document

from app.core.config import settings


def render_docx_placeholders(doc: Document, replacements: dict[str, str]) -> None:
    """Replace placeholders in a DOCX document.

    Walks paragraphs in the document body and all table cells, applying
    replacements using run-coordinate mapping so formatting is preserved.
    """
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
    """Replace placeholders in a single paragraph using run coordinate mapping.

    All replacements are applied in a single pass (right-to-left) so run indices stay valid.
    This is the same algorithm used in order_document_service.
    """
    if not paragraph.runs:
        return

    full_text = paragraph.text

    # Collect all occurrences of all keys
    occurrences: list[tuple[int, str, str]] = []  # (start_pos, key, value)
    for key, value in replacements.items():
        if key not in full_text:
            continue
        key_len = len(key)
        for i in range(len(full_text)):
            if full_text.startswith(key, i):
                occurrences.append((i, key, value))

    if not occurrences:
        return

    # Sort by position descending (right-to-left) so earlier indices stay valid
    occurrences.sort(key=lambda x: x[0], reverse=True)

    # Build coordinate map once
    p_map = []
    for run_idx, run in enumerate(paragraph.runs):
        for char_idx, char in enumerate(run.text):
            p_map.append({"run": run_idx, "char": char_idx})

    for start_pos, key, value in occurrences:
        key_len = len(key)
        if start_pos + key_len > len(p_map):
            continue
        key_map = p_map[start_pos : start_pos + key_len]
        _replace_in_runs(paragraph.runs, key_map, value)


def _replace_in_runs(runs: list[Any], key_map: list[dict], value: str) -> None:
    """Apply replacement to the specific runs/characters identified by key_map."""
    for i, position in enumerate(reversed(key_map), start=1):
        run_idx = position["run"]
        char_idx = position["char"]
        run = runs[run_idx]
        chars = list(run.text)

        if i < len(key_map):
            chars.pop(char_idx)
        else:
            chars[char_idx] = value

        run.text = "".join(chars)


async def load_template_or_create_blank(template_path: Path) -> Document:
    """Load a DOCX template if it exists, otherwise create a blank document."""
    if template_path.exists():
        return await asyncio.wait_for(
            asyncio.to_thread(Document, str(template_path)),
            timeout=settings.DOCUMENT_GENERATION_TIMEOUT,
        )
    return Document()


def build_basic_doc_replacements(
    title: str = "",
    number: str = "",
    date_str: str = "",
    extra: dict[str, str] | None = None,
) -> dict[str, str]:
    """Build a basic set of replacements common to all document types.

    This provides a safe fallback set of placeholders for quick-draft mode.
    """
    replacements = {
        "{doc_title}": title,
        "{doc_number}": number,
        "{doc_date}": date_str,
    }
    if extra:
        for key, value in extra.items():
            replacements[f"{{{key}}}"] = str(value)
    return replacements
