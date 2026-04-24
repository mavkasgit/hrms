"""Генерация универсального fallback-шаблона __missing__.docx.

Используется, когда у типа приказа нет загруженного шаблона.
Содержит:
- уведомление о том, что шаблон не загружен,
- базовые placeholders для демонстрации работы подстановки.
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
except ImportError as exc:
    print("ERROR: python-docx не установлен. Установите: pip install python-docx")
    raise SystemExit(1) from exc


TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "data" / "templates"
OUTPUT_PATH = TEMPLATES_DIR / "__missing__.docx"


def main() -> int:
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Заголовок
    heading = doc.add_heading("Приказ №{order_number}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Уведомление
    warning_para = doc.add_paragraph()
    warning_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warning_para.add_run(
        "ВНИМАНИЕ! Документ сгенерирован без шаблона. "
        "Обязательно загрузите шаблон для этого типа приказа в настройках."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.size = Pt(12)

    # Путь к настройкам
    hint_para = doc.add_paragraph()
    hint_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = hint_para.add_run("Настройки шаблонов: /templates")
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    run2.font.size = Pt(11)

    doc.add_paragraph()  # отступ

    # Блок с основными placeholders
    doc.add_paragraph("Основные данные приказа:").runs[0].bold = True

    fields = [
        ("Тип приказа", "{order_type_name}"),
        ("Дата приказа", "{order_date}"),
        ("Сотрудник (ФИО)", "{full_name}"),
        ("Сотрудник (кратко)", "{short_name}"),
        ("Должность", "{position}"),
        ("Подразделение", "{department}"),
        ("Табельный номер", "{tab_number}"),
        ("Дата приема", "{hire_date}"),
        ("Примечания", "{notes}"),
    ]

    for label, placeholder in fields:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        p.add_run(placeholder)

    doc.add_paragraph()  # отступ

    # Блок с дополнительными placeholders
    doc.add_paragraph("Дополнительные placeholders (при наличии):").runs[0].bold = True
    extras = [
        "{full_name_upper}", "{last_name_upper}", "{initials_before}",
        "{position_cap}", "{order_type_code}", "{oznak_gender}",
    ]
    for ph in extras:
        doc.add_paragraph(ph, style="List Bullet")

    doc.save(str(OUTPUT_PATH))
    print(f"OK: {OUTPUT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
