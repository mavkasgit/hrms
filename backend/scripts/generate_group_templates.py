"""Generate sample DOCX templates for group orders with proper placeholder markers."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

TEMPLATES_DIR = Path(__file__).parent.parent / "data" / "templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


def add_paragraph_styled(doc, text, bold=False, size=12, alignment=None, space_after=6):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_placeholder_line(doc, label, placeholder, bold_label=True):
    """Add a line with label and placeholder."""
    p = doc.add_paragraph()
    if bold_label:
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(12)
    run_value = p.add_run(f" {placeholder}")
    run_value.font.size = Pt(12)
    run_value.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)  # Blue for visibility
    return p


def create_vacation_unpaid_group_template():
    """Create sample template for vacation_unpaid_group orders."""
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # Header
    add_paragraph_styled(doc, "ПРИКАЗ", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_styled(doc, "о предоставлении отпуска без сохранения заработной платы", bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Order info
    add_placeholder_line(doc, "Номер приказа: ", "{order_number}")
    add_placeholder_line(doc, "Дата приказа: ", "{order_date}")
    add_paragraph_styled(doc, "", space_after=6)

    # Main text
    p = doc.add_paragraph()
    p.add_run("ПРЕДСТАВИТЬ ").bold = True
    p.add_run("отпуск без сохранения заработной платы ")

    # Employee block start marker (MUST be on its own paragraph)
    add_paragraph_styled(doc, "{employees_block_start}", size=10, space_after=0)

    # Employee block template (will be repeated for each employee)
    p = doc.add_paragraph()
    p.add_run("{index}. ")
    run_name = p.add_run("{full_name_last_caps}")
    run_name.bold = True
    p.add_run(", ")
    p.add_run("{position}")

    p2 = doc.add_paragraph()
    p2.add_run("   на ")
    run_days = p2.add_run("{vacation_days}")
    run_days.bold = True
    p2.add_run(" календарных дней")
    p2.add_run(" с ")
    run_start = p2.add_run("{vacation_start}")
    run_start.bold = True
    p2.add_run(" по ")
    run_end = p2.add_run("{vacation_end}")
    run_end.bold = True

    p3 = doc.add_paragraph()
    p3.add_run("   Заявление от: ")
    p3.add_run("{application_date}")

    # Employee block end marker (MUST be on its own paragraph)
    add_paragraph_styled(doc, "{employees_block_end}", size=10, space_after=12)

    # Footer section
    add_paragraph_styled(doc, "Основание: заявления работников", bold=True, space_after=12)

    # Application block (for showing employee signatures)
    add_paragraph_styled(doc, "С заявлениями ознакомлены:", bold=True, space_after=6)

    add_paragraph_styled(doc, "{applications_block_start}", size=10, space_after=0)

    p = doc.add_paragraph()
    p.add_run("{index}. ")
    p.add_run("{short_name}")
    p.add_run(" _______________   «{application_date}»")

    add_paragraph_styled(doc, "{applications_block_end}", size=10, space_after=12)

    # Signature section
    add_paragraph_styled(doc, "", space_after=24)
    add_placeholder_line(doc, "Руководитель: ", "{initials_before}")
    add_paragraph_styled(doc, "", space_after=12)
    add_placeholder_line(doc, "Документ составлен и проверен, сотрудник ", "{oznak_gender}")

    # Save
    output_path = TEMPLATES_DIR / "template__order__vacation_unpaid_group.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


def create_weekend_call_group_template():
    """Create sample template for weekend_call_group orders."""
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # Header
    add_paragraph_styled(doc, "ПРИКАЗ", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_styled(doc, "о привлечении к работе в выходной день", bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Order info
    add_placeholder_line(doc, "Номер приказа: ", "{order_number}")
    add_placeholder_line(doc, "Дата приказа: ", "{order_date}")
    add_paragraph_styled(doc, "", space_after=6)

    # Main text
    p = doc.add_paragraph()
    run1 = p.add_run("ПРИВЛЕЧЬ ")
    run1.bold = True
    p.add_run("к работе в выходной день ")
    run_period = p.add_run("с {call_date_start} по {call_date_end}")
    run_period.bold = True

    # Employee block start marker (MUST be on its own paragraph)
    add_paragraph_styled(doc, "{employees_block_start}", size=10, space_after=0)

    # Employee block template (will be repeated for each employee)
    p = doc.add_paragraph()
    p.add_run("{index}. ")
    run_name = p.add_run("{full_name_last_caps}")
    run_name.bold = True
    p.add_run(", ")
    p.add_run("{position}")

    p2 = doc.add_paragraph()
    p2.add_run("   на ")
    run_days = p2.add_run("{vacation_days}")
    run_days.bold = True
    p2.add_run(" календарных дней")
    p2.add_run(" в период ")
    run_start = p2.add_run("с {call_date_start}")
    run_start.bold = True
    p2.add_run(" по ")
    run_end = p2.add_run("{call_date_end}")
    run_end.bold = True

    p3 = doc.add_paragraph()
    p3.add_run("   Заявление от: ")
    p3.add_run("{application_date}")

    # Employee block end marker (MUST be on its own paragraph)
    add_paragraph_styled(doc, "{employees_block_end}", size=10, space_after=12)

    # Footer section
    add_paragraph_styled(doc, "Основание: заявления работников", bold=True, space_after=12)

    # Application block (for showing employee signatures)
    add_paragraph_styled(doc, "С заявлениями ознакомлены:", bold=True, space_after=6)

    add_paragraph_styled(doc, "{applications_block_start}", size=10, space_after=0)

    p = doc.add_paragraph()
    p.add_run("{index}. ")
    p.add_run("{short_name}")
    p.add_run(" _______________   «{application_date}»")

    add_paragraph_styled(doc, "{applications_block_end}", size=10, space_after=12)

    # Signature section
    add_paragraph_styled(doc, "", space_after=24)
    add_placeholder_line(doc, "Руководитель: ", "{initials_before}")
    add_paragraph_styled(doc, "", space_after=12)
    add_placeholder_line(doc, "Документ составлен и проверен, сотрудник ", "{oznak_gender}")

    # Save
    output_path = TEMPLATES_DIR / "template__order__weekend_call_group.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Generating sample group order templates...")
    create_vacation_unpaid_group_template()
    create_weekend_call_group_template()
    print("Done!")
