from datetime import date
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document

from app.core.exceptions import HRMSException
from app.schemas.order import OrderCreate
from app.services.order_service import order_service


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _paragraph_texts(docx_payload: bytes) -> list[str]:
    document = Document(BytesIO(docx_payload))
    return [paragraph.text for paragraph in document.paragraphs]


def test_extract_paragraph_texts_from_html_normalizes_markup():
    html = "<p><strong>Первый</strong> абзац</p><p>Второй<br>ряд</p><div>Третий</div>"

    texts = order_service._extract_paragraph_texts_from_html(html)

    assert texts == ["Первый абзац", "Второй\nряд", "Третий"]


def test_apply_edited_html_to_docx_replaces_only_changed_paragraphs():
    original_docx = _docx_bytes("Заголовок", "Старый текст", "Подпись")
    original_html = "<p>Заголовок</p><p>Старый текст</p><p>Подпись</p>"
    edited_html = "<p>Заголовок</p><p>Новый текст</p><p>Подпись</p>"

    updated_docx = order_service._apply_edited_html_to_docx(original_docx, original_html, edited_html)

    assert _paragraph_texts(updated_docx) == ["Заголовок", "Новый текст", "Подпись"]


def test_apply_edited_html_to_docx_updates_paragraphs_inside_tables():
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "Текст в таблице"
    stream = BytesIO()
    document.save(stream)

    original_html = "<p>Текст в таблице</p>"
    edited_html = "<p>Изменено в таблице</p>"

    updated_docx = order_service._apply_edited_html_to_docx(stream.getvalue(), original_html, edited_html)
    updated_document = Document(BytesIO(updated_docx))

    assert updated_document.tables[0].cell(0, 0).paragraphs[0].text == "Изменено в таблице"


def test_highlight_placeholder_values_wraps_inserted_text():
    html = "<p>Сотрудник: Иванов Иван Иванович</p><p>Дата: 03.04.2026</p>"
    replacements = {
        "{full_name}": "Иванов Иван Иванович",
        "{order_date}": "03.04.2026",
    }

    highlighted = order_service._highlight_placeholder_values(html, replacements)

    assert 'data-placeholder-key="full_name"' in highlighted
    assert 'data-placeholder-key="order_date"' in highlighted
    assert "<mark" in highlighted


def test_highlight_placeholder_values_does_not_nest_marks_for_overlapping_values():
    html = "<p>Номер: TEST-001</p><p>Табельный: 001</p>"
    replacements = {
        "{order_number}": "TEST-001",
        "{tab_number}": "001",
    }

    highlighted = order_service._highlight_placeholder_values(html, replacements)

    assert highlighted.count('data-placeholder-key="order_number"') == 1
    assert highlighted.count('data-placeholder-key="tab_number"') == 1
    assert '<mark class="order-placeholder" data-placeholder-key="order_number">TEST-001</mark>' in highlighted
    assert 'TEST-<mark' not in highlighted


def test_style_missing_template_warning_html_marks_warning_paragraph():
    html = "<h1>Приказ №1</h1><p>ВНИМАНИЕ: документ сгенерирован без шаблона.</p><p>Тип: Тест</p>"

    styled = order_service._style_missing_template_warning_html(html)

    assert '<p class="missing-template-warning"><strong>ВНИМАНИЕ: документ сгенерирован без шаблона.</strong></p>' in styled


@pytest.mark.asyncio
async def test_generate_order_preview_uses_missing_template_when_no_template(
    monkeypatch,
):
    employee = SimpleNamespace(
        id=1,
        name="Иванов Иван Иванович",
        gender="male",
        tab_number="001",
        department=SimpleNamespace(name="Отдел кадров"),
        position=SimpleNamespace(name="Инспектор"),
        hire_date=date(2024, 1, 1),
        contract_start=date(2024, 1, 1),
    )
    order_type = SimpleNamespace(
        id=2,
        is_active=True,
        code="test_no_template",
        name="Тест без шаблона",
        template_filename=None,
        filename_pattern=None,
    )

    async def ensure_defaults(_db):
        return []

    async def get_employee(_db, _employee_id):
        return employee

    async def get_order_type(_db, _order_type_id):
        return order_type

    monkeypatch.setattr(order_service, "ensure_default_order_types", ensure_defaults)
    monkeypatch.setattr(order_service.employee_repo, "get_by_id", get_employee)
    monkeypatch.setattr(order_service.order_type_repo, "get_by_id", get_order_type)

    data = OrderCreate(
        employee_id=employee.id,
        order_type_id=order_type.id,
        order_date=date(2024, 1, 15),
        order_number="TEST-001",
    )

    preview = await order_service.generate_order_preview(None, data)

    assert preview["preview_id"]
    assert "TEST-001" in preview["html"]
    assert 'data-placeholder-key="full_name"' in preview["html"]
    assert 'data-placeholder-key="order_number"' in preview["html"]
    assert 'class="missing-template-warning"' in preview["html"]
    assert "шаблон" in preview["html"].lower()
