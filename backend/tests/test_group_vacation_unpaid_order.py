from datetime import date

import pytest
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.models.order import Order
from app.models.order_employee import OrderEmployee
from app.models.vacation import Vacation
from app.schemas.order import VacationUnpaidGroupOrderCreate, VacationUnpaidGroupEmployeeCreate
from app.services.order_service import order_service

pytestmark = pytest.mark.asyncio(loop_scope="module")


async def test_create_vacation_unpaid_group_order(db_session, create_employee, create_position, create_department):
    """Создание группового приказа: проверка Order, OrderEmployee, Vacation и сериализации."""
    await order_service.ensure_default_order_types(db_session)

    dept = await create_department(name="Тестовый отдел")
    pos = await create_position(name="Тестовая должность")

    emp1 = await create_employee(name="Иванов Иван Иванович", department=dept, position=pos)
    emp2 = await create_employee(name="Петров Пётр Петрович", department=dept, position=pos)
    emp3 = await create_employee(name="Сидорова Анна Сергеевна", department=dept, position=pos)

    payload = VacationUnpaidGroupOrderCreate(
        order_date=date(2026, 5, 10),
        order_number="42-Т",
        vacation_start=date(2026, 5, 15),
        employees=[
            VacationUnpaidGroupEmployeeCreate(employee_id=emp1.id, vacation_days=5),
            VacationUnpaidGroupEmployeeCreate(employee_id=emp2.id, vacation_days=10),
            VacationUnpaidGroupEmployeeCreate(employee_id=emp3.id, vacation_days=1),
        ],
    )

    order = await order_service.create_vacation_unpaid_group_order(db_session, payload)

    # Проверка Order
    assert order is not None
    assert order.order_number == "42-Т"
    assert order.order_date == date(2026, 5, 10)
    assert order.employee_id is None
    assert order.is_group is True

    # Проверка через прямую выборку из БД
    result = await db_session.execute(
        select(OrderEmployee)
        .where(OrderEmployee.order_id == order.id)
        .order_by(OrderEmployee.employee_id)
    )
    order_employees = list(result.scalars().all())
    assert len(order_employees) == 3

    # Проверяем каждого сотрудника
    emp_map = {oe.employee_id: oe for oe in order_employees}

    assert emp_map[emp1.id].vacation_start == date(2026, 5, 15)
    assert emp_map[emp1.id].vacation_end == date(2026, 5, 19)  # 15 + 5 - 1 = 19
    assert emp_map[emp1.id].vacation_days == 5

    assert emp_map[emp2.id].vacation_start == date(2026, 5, 15)
    assert emp_map[emp2.id].vacation_end == date(2026, 5, 24)  # 15 + 10 - 1 = 24
    assert emp_map[emp2.id].vacation_days == 10

    assert emp_map[emp3.id].vacation_start == date(2026, 5, 15)
    assert emp_map[emp3.id].vacation_end == date(2026, 5, 15)  # 15 + 1 - 1 = 15
    assert emp_map[emp3.id].vacation_days == 1

    # Проверка Vacation records
    vac_result = await db_session.execute(
        select(Vacation).where(Vacation.order_id == order.id).order_by(Vacation.employee_id)
    )
    vacations = list(vac_result.scalars().all())
    assert len(vacations) == 3

    vac_map = {v.employee_id: v for v in vacations}
    assert vac_map[emp1.id].start_date == date(2026, 5, 15)
    assert vac_map[emp1.id].end_date == date(2026, 5, 19)
    assert vac_map[emp1.id].days_count == 5

    assert vac_map[emp2.id].start_date == date(2026, 5, 15)
    assert vac_map[emp2.id].end_date == date(2026, 5, 24)
    assert vac_map[emp2.id].days_count == 10

    assert vac_map[emp3.id].start_date == date(2026, 5, 15)
    assert vac_map[emp3.id].end_date == date(2026, 5, 15)
    assert vac_map[emp3.id].days_count == 1

    # Проверка сериализации (имя, код типа, group_employees)
    serialized = order_service._serialize_order(order)
    assert serialized["is_group"] is True
    assert serialized["group_employee_count"] == 3
    assert serialized["order_type_code"] == "vacation_unpaid_group"

    group_emps = serialized["group_employees"]
    assert len(group_emps) == 3

    ge_map = {ge["employee_id"]: ge for ge in group_emps}
    assert ge_map[emp1.id]["employee_full_name"] == "Иванов Иван Иванович"
    assert ge_map[emp1.id]["vacation_days"] == 5
    assert ge_map[emp1.id]["vacation_start"] == "2026-05-15"
    assert ge_map[emp1.id]["vacation_end"] == "2026-05-19"

    assert ge_map[emp2.id]["employee_full_name"] == "Петров Пётр Петрович"
    assert ge_map[emp2.id]["vacation_days"] == 10
    assert ge_map[emp2.id]["vacation_start"] == "2026-05-15"
    assert ge_map[emp2.id]["vacation_end"] == "2026-05-24"

    assert ge_map[emp3.id]["employee_full_name"] == "Сидорова Анна Сергеевна"
    assert ge_map[emp3.id]["vacation_days"] == 1
    assert ge_map[emp3.id]["vacation_start"] == "2026-05-15"
    assert ge_map[emp3.id]["vacation_end"] == "2026-05-15"

    # Проверка DOCX файла
    from app.core.paths import storage_path
    assert order.file_path is not None
    docx_path = storage_path(order.file_path, "ORDERS_PATH")
    assert docx_path.exists(), f"DOCX файл не найден: {docx_path}"

    # Проверяем что DOCX содержит ФИО и дни всех сотрудников
    from docx import Document
    doc = Document(docx_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    assert "Иванов Иван Иванович" in full_text
    assert "Петров Пётр Петрович" in full_text
    assert "Сидорова Анна Сергеевна" in full_text
