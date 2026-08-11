from datetime import date
from pathlib import Path
from types import SimpleNamespace
from typing import Any, cast
from unittest.mock import AsyncMock, MagicMock
import os
import time

import pytest
from docx import Document
from jose import jwt
from fastapi import Request
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.exceptions import HRMSException
from app.api.orders import print_order_pdf
from app.models.employee import Employee
from app.models.order_type import OrderType
from app.schemas.order import OrderCreate
from app.services.onlyoffice_service import OnlyOfficeService
from app.services.order_draft_service import OrderDraftService
from app.services.order_print_service import OrderPrintService, order_print_service
from app.services.order_service import order_service


def _emp(**kw: Any) -> Employee:
    return cast(Employee, SimpleNamespace(**kw))


def _ot(**kw: Any) -> OrderType:
    return cast(OrderType, SimpleNamespace(**kw))


def _db() -> AsyncSession:
    return cast(AsyncSession, object())


def test_onlyoffice_config_print_allowed(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    docx_path = tmp_path / "order.docx"
    docx_path.write_bytes(b"PK")

    config = OnlyOfficeService().build_config(
        doc_type="order",
        doc_id=1,
        file_path=docx_path,
        title="order.docx",
        callback_url="http://app/api/orders/1/onlyoffice/callback",
        file_url="http://app/api/orders/1/onlyoffice/file",
        allow_print=True,
    )

    assert config["document"]["permissions"]["print"] is True


def test_onlyoffice_config_print_disabled(monkeypatch, tmp_path):
    """Draft config should have print: False, and the JWT token should reflect it."""
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    docx_path = tmp_path / "draft.docx"
    docx_path.write_bytes(b"PK")

    config = OnlyOfficeService().build_config(
        doc_type="draft",
        doc_id="test-draft-id",
        file_path=docx_path,
        title="draft.docx",
        callback_url="http://app/api/orders/drafts/test-draft-id/onlyoffice/callback",
        file_url="http://app/api/orders/drafts/test-draft-id/onlyoffice/file",
        allow_print=False,
    )

    assert config["document"]["permissions"]["print"] is False
    decoded = jwt.decode(config["token"], "test-secret", algorithms=["HS256"])
    assert decoded["document"]["permissions"]["print"] is False


def test_onlyoffice_config_contains_signed_token(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    docx_path = tmp_path / "order.docx"
    docx_path.write_bytes(b"PK")

    config = OnlyOfficeService().build_config(
        doc_type="order",
        doc_id=1,
        file_path=docx_path,
        title="order.docx",
        callback_url="http://app/api/orders/1/onlyoffice/callback",
        file_url="http://app/api/orders/1/onlyoffice/file",
    )

    assert config["document"]["fileType"] == "docx"
    assert config["document"]["key"].startswith("order-1-")
    decoded = jwt.decode(config["token"], "test-secret", algorithms=["HS256"])
    assert decoded["document"]["url"] == "http://app/api/orders/1/onlyoffice/file"


def test_onlyoffice_config_includes_data_array(monkeypatch, tmp_path):
    """document.data (массив полей) попадает в конфиг и в подписанный JWT."""
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    docx_path = tmp_path / "order.docx"
    docx_path.write_bytes(b"PK")

    data = [{"key": "order_number", "value": "42-К"}, {"key": "order_date", "value": "2026-08-06"}]

    config = OnlyOfficeService().build_config(
        doc_type="draft",
        doc_id="draft-id",
        file_path=docx_path,
        title="draft.docx",
        callback_url="http://app/api/orders/drafts/draft-id/onlyoffice/callback",
        file_url="http://app/api/orders/drafts/draft-id/file",
        data=data,
    )

    assert config["document"]["data"] == data
    decoded = jwt.decode(config["token"], "test-secret", algorithms=["HS256"])
    assert decoded["document"]["data"] == data


def test_onlyoffice_config_omits_empty_data_array(monkeypatch, tmp_path):
    """Пустой data-массив не попадает в конфиг."""
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    docx_path = tmp_path / "order.docx"
    docx_path.write_bytes(b"PK")

    config = OnlyOfficeService().build_config(
        doc_type="order",
        doc_id=1,
        file_path=docx_path,
        title="order.docx",
        callback_url="http://app/api/orders/1/onlyoffice/callback",
        file_url="http://app/api/orders/1/onlyoffice/file",
        data=[],
    )

    assert "data" not in config["document"]


def test_onlyoffice_config_includes_user(monkeypatch, tmp_path):
    """editorConfig.user (id/name) попадает в конфиг и в подписанный JWT (ADR-0010)."""
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    docx_path = tmp_path / "order.docx"
    docx_path.write_bytes(b"PK")

    config = OnlyOfficeService().build_config(
        doc_type="order",
        doc_id=1,
        file_path=docx_path,
        title="order.docx",
        callback_url="http://app/api/orders/1/onlyoffice/callback",
        file_url="http://app/api/orders/1/onlyoffice/file",
        user={"id": "ivanov", "name": "Иванов Иван"},
    )

    assert config["editorConfig"]["user"] == {"id": "ivanov", "name": "Иванов Иван"}
    decoded = jwt.decode(config["token"], "test-secret", algorithms=["HS256"])
    assert decoded["editorConfig"]["user"]["id"] == "ivanov"
    assert decoded["editorConfig"]["user"]["name"] == "Иванов Иван"


def test_onlyoffice_config_omits_user_when_not_passed(monkeypatch, tmp_path):
    """Без user параметра ключ editorConfig.user не появляется."""
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    docx_path = tmp_path / "order.docx"
    docx_path.write_bytes(b"PK")

    config = OnlyOfficeService().build_config(
        doc_type="order",
        doc_id=1,
        file_path=docx_path,
        title="order.docx",
        callback_url="http://app/api/orders/1/onlyoffice/callback",
        file_url="http://app/api/orders/1/onlyoffice/file",
    )

    assert "user" not in config["editorConfig"]


def test_onlyoffice_config_disables_features_tips(monkeypatch, tmp_path):
    """What's new попап отключается features.featuresTips=false (ADR-0010)."""
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    docx_path = tmp_path / "order.docx"
    docx_path.write_bytes(b"PK")

    config = OnlyOfficeService().build_config(
        doc_type="order",
        doc_id=1,
        file_path=docx_path,
        title="order.docx",
        callback_url="http://app/api/orders/1/onlyoffice/callback",
        file_url="http://app/api/orders/1/onlyoffice/file",
    )

    assert config["editorConfig"]["customization"]["features"] == {"featuresTips": False}
    decoded = jwt.decode(config["token"], "test-secret", algorithms=["HS256"])
    assert decoded["editorConfig"]["customization"]["features"]["featuresTips"] is False


def test_onlyoffice_form_data_builds_array(monkeypatch):
    """build_form_data: None и пустые строки пропускаются, list/dict отбрасываются."""
    from app.services.onlyoffice_form_data import build_form_data

    data = build_form_data({
        "order_number": "42-К",
        "order_date": "2026-08-06",
        "vacation_days": 14,
        "notes": "",
        "empty": None,
        "employees": [{"employee_id": 1}],
    })

    assert data == [
        {"key": "order_number", "value": "42-К"},
        {"key": "order_date", "value": "2026-08-06"},
        {"key": "vacation_days", "value": "14"},
    ]


def test_onlyoffice_draft_form_data_from_metadata(monkeypatch):
    """draft_form_data собирает массив из payload метаданных черновика."""
    from app.services.onlyoffice_form_data import draft_form_data

    data = draft_form_data({
        "kind": "group_order",
        "payload": {
            "order_number": "10-К",
            "order_date": "2026-08-01",
            "vacation_start": "2026-09-01",
            "mode": "single",
            "call_date": "2026-09-05",
            "extra_fields": None,
            "employees": [{"employee_id": 1}, {"employee_id": 2}],
        },
    })

    assert data == [
        {"key": "number", "value": "10-К"},
        {"key": "date", "value": "2026-08-01"},
        {"key": "vacation_start", "value": "2026-09-01"},
        {"key": "mode", "value": "single"},
        {"key": "call_date", "value": "2026-09-05"},
    ]


def test_onlyoffice_callback_token_validation(monkeypatch):
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    service = OnlyOfficeService()
    token = jwt.encode({"status": 2}, "test-secret", algorithm="HS256")

    assert service.validate_callback_token(token) is True
    assert service.validate_callback_token("bad-token") is False


def test_onlyoffice_download_url_uses_internal_url(monkeypatch):
    monkeypatch.setattr(settings, "ONLYOFFICE_PUBLIC_URL", "http://localhost:8085")
    monkeypatch.setattr(settings, "ONLYOFFICE_INTERNAL_URL", "http://onlyoffice:80")

    normalized = OnlyOfficeService()._normalize_download_url("http://localhost:8085/cache/doc.docx")

    assert normalized == "http://onlyoffice:80/cache/doc.docx"


@pytest.mark.asyncio
async def test_replace_docx_atomically(tmp_path):
    target = tmp_path / "target.docx"
    temp = tmp_path / "temp.docx"
    target.write_bytes(b"old")
    temp.write_bytes(b"new")

    await OnlyOfficeService()._replace_docx_atomically(target, temp)

    assert target.read_bytes() == b"new"
    assert not temp.exists()


@pytest.mark.asyncio
async def test_order_draft_service_creates_docx(monkeypatch, tmp_path):
    service = OrderDraftService()
    service._drafts_dir = tmp_path / ".drafts"

    document = Document()
    document.add_paragraph("Черновик")

    async def fake_build_document(*_args, **_kwargs):
        return document, {"{order_number}": "1"}

    monkeypatch.setattr("app.services.order_draft_service._build_document", fake_build_document)
    monkeypatch.setattr("app.services.order_draft_service._build_filename", lambda *_args: "order.docx")

    draft = await service.create_draft(
        OrderCreate(employee_id=1, order_type_id=2, order_date=date.today(), order_number="1"),
        _emp(name="Иванов Иван Иванович"),
        _ot(code="test", name="Тест"),
    )

    assert draft["draft_id"]
    assert Path(draft["file_path"]).exists()
    # Метаданные сохранены рядом с docx (#29)
    metadata_path = service.get_metadata_path(draft["draft_id"])
    assert metadata_path.exists()


@pytest.mark.asyncio
async def test_order_draft_metadata_content(monkeypatch, tmp_path):
    """#29: метаданные содержат полный payload, создателя, время, статус и версию схемы."""
    import json

    service = OrderDraftService()
    service._drafts_dir = tmp_path / ".drafts"

    document = Document()
    document.add_paragraph("Черновик")

    async def fake_build_document(*_args, **_kwargs):
        return document, {"{order_number}": "42-к"}

    monkeypatch.setattr("app.services.order_draft_service._build_document", fake_build_document)
    monkeypatch.setattr("app.services.order_draft_service._build_filename", lambda *_args: "order.docx")

    draft = await service.create_draft(
        OrderCreate(
            employee_id=7,
            order_type_id=3,
            order_date=date(2026, 8, 1),
            order_number="42-к",
            notes="тест",
            extra_fields={"vacation_start": "2026-08-10"},
        ),
        _emp(name="Петров П.П."),
        _ot(code="vacation_paid", name="Отпуск"),
        user_id="admin-user",
    )

    metadata = json.loads(service.get_metadata_path(draft["draft_id"]).read_text(encoding="utf-8"))
    assert metadata["draft_id"] == draft["draft_id"]
    assert metadata["kind"] == "single_order"
    assert metadata["order_type_code"] == "vacation_paid"
    assert metadata["payload"]["employee_id"] == 7
    assert metadata["payload"]["order_type_id"] == 3
    assert metadata["payload"]["order_date"] == "2026-08-01"
    assert metadata["payload"]["order_number"] == "42-к"
    assert metadata["payload"]["notes"] == "тест"
    assert metadata["payload"]["extra_fields"] == {"vacation_start": "2026-08-10"}
    assert metadata["created_by"] == "admin-user"
    assert metadata["status"] == "draft"
    assert metadata["schema_version"] == 1
    assert "created_at" in metadata


@pytest.mark.asyncio
async def test_order_draft_metadata_failure_rolls_back_docx(monkeypatch, tmp_path):
    """#29: при ошибке сохранения метаданных черновик откатывается целиком."""
    service = OrderDraftService()
    service._drafts_dir = tmp_path / ".drafts"

    document = Document()
    document.add_paragraph("Черновик")

    async def fake_build_document(*_args, **_kwargs):
        return document, {"{order_number}": "1"}

    monkeypatch.setattr("app.services.order_draft_service._build_document", fake_build_document)
    monkeypatch.setattr("app.services.order_draft_service._build_filename", lambda *_args: "order.docx")

    # Ломаем сохранение метаданных
    def broken_save_metadata(draft_id, metadata):
        raise OSError("disk full")

    monkeypatch.setattr(service, "save_draft_metadata", broken_save_metadata)

    with pytest.raises(OSError, match="disk full"):
        await service.create_draft(
            OrderCreate(employee_id=1, order_type_id=2, order_date=date.today(), order_number="1"),
            _emp(name="Иванов"),
            _ot(code="test", name="Тест"),
        )

    # Docx не должен остаться
    drafts_dir = tmp_path / ".drafts"
    docx_files = list(drafts_dir.glob("*.docx")) if drafts_dir.exists() else []
    assert docx_files == []


def test_order_draft_service_rejects_unknown_draft(tmp_path):
    service = OrderDraftService()
    service._drafts_dir = tmp_path / ".drafts"

    with pytest.raises(HRMSException) as exc_info:
        service.get_draft_path("not-a-draft")

    assert exc_info.value.status_code == 404


@pytest.mark.asyncio
async def test_generate_document_from_draft_keeps_draft(monkeypatch, tmp_path):
    """generate_document copies draft to permanent storage but does NOT delete the draft."""
    from app.services.order_document_service import generate_document

    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))

    draft_service = OrderDraftService()
    draft_service._drafts_dir = tmp_path / ".drafts"
    draft_service.ensure_drafts_dir()
    draft_id = "12345678-1234-1234-1234-123456789abc"
    draft_path = draft_service._drafts_dir / f"{draft_id}_order.docx"
    draft_path.write_bytes(b"draft")

    monkeypatch.setattr("app.services.order_draft_service.order_draft_service", draft_service)
    monkeypatch.setattr("app.services.order_document_service._build_document", AsyncMock(return_value=(Document(), {"{order_number}": "1"})))
    monkeypatch.setattr("app.services.order_document_service._build_filename", lambda *_args: "final.docx")

    result = await generate_document(
        "1",
        OrderCreate(
            employee_id=1,
            order_type_id=2,
            order_date=date.today(),
            order_number="1",
            draft_id=draft_id,
        ),
        _emp(name="Иванов Иван Иванович"),
        _ot(code="test", name="Тест", filename_pattern="final.docx"),
        tmp_path,
    )

    assert (tmp_path / result[0]).read_bytes() == b"draft"
    # Draft must remain until create_order succeeds (retry-safe).
    assert draft_path.exists()


@pytest.mark.asyncio
async def test_create_order_deletes_draft_after_success(monkeypatch, tmp_path):
    """Draft is deleted only after _do_create_order succeeds (group-commit pattern)."""
    from app.services.order_draft_service import order_draft_service
    from app.services.order_service import OrderService

    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    order_draft_service._drafts_dir = tmp_path / ".drafts"
    order_draft_service.ensure_drafts_dir()

    draft_id = "aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee"
    draft_path = order_draft_service._drafts_dir / f"{draft_id}_order.docx"
    draft_path.write_bytes(b"draft-content")

    svc = OrderService()
    fake_order = SimpleNamespace(id=42)

    async def fake_ensure(db):
        return []

    async def fake_do_create(db, data):
        assert data.draft_id == draft_id
        assert draft_path.exists(), "draft must still exist during create"
        return fake_order

    monkeypatch.setattr(svc, "ensure_default_order_types", fake_ensure)
    monkeypatch.setattr(svc, "_do_create_order", fake_do_create)

    db = MagicMock()
    db.in_transaction.return_value = True

    order = await svc.create_order(
        db,
        OrderCreate(
            employee_id=None,
            order_type_id=1,
            order_date=date.today(),
            order_number="DRAFT-OK-1",
            draft_id=draft_id,
        ),
    )

    assert order.id == 42
    assert not draft_path.exists()


@pytest.mark.asyncio
async def test_create_order_keeps_draft_when_create_fails(monkeypatch, tmp_path):
    """If create fails after draft copy, draft remains and permanent file is cleaned up."""
    from app.services import order_service as order_service_module
    from app.services.order_draft_service import order_draft_service
    from app.services.order_service import OrderService

    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    order_draft_service._drafts_dir = tmp_path / ".drafts"
    order_draft_service.ensure_drafts_dir()

    draft_id = "11111111-2222-3333-4444-555555555555"
    draft_path = order_draft_service._drafts_dir / f"{draft_id}_order.docx"
    draft_path.write_bytes(b"draft-retry")

    permanent_written: list[Path] = []

    async def fake_generate(order_number, data, employee, order_type, year_dir_arg):
        year_dir_arg.mkdir(parents=True, exist_ok=True)
        dest = year_dir_arg / f"{order_type.code}_{order_number}.docx"
        dest.write_bytes(b"copied-from-draft")
        permanent_written.append(dest)
        return f"{data.order_date.year}/{dest.name}", dest.name

    async def failing_finish(*_args, **_kwargs):
        raise RuntimeError("simulated DB failure")

    svc = OrderService()
    order_type = SimpleNamespace(id=1, code="general_order", name="Общий", is_active=True)

    async def fake_ensure(db):
        return [order_type]

    async def fake_get_type(db, type_id):
        return order_type

    monkeypatch.setattr(order_service_module, "generate_document", fake_generate)
    monkeypatch.setattr(svc, "ensure_default_order_types", fake_ensure)
    monkeypatch.setattr(svc.order_type_repo, "get_by_id", fake_get_type)
    monkeypatch.setattr(svc, "_finish_create_order", failing_finish)

    db = MagicMock()
    db.in_transaction.return_value = True

    with pytest.raises(RuntimeError, match="simulated DB failure"):
        await svc.create_order(
            db,
            OrderCreate(
                employee_id=None,
                order_type_id=1,
                order_date=date.today(),
                order_number="DRAFT-FAIL-1",
                draft_id=draft_id,
            ),
        )

    assert draft_path.exists(), "draft must remain for retry after failed create"
    for p in permanent_written:
        assert not p.exists(), f"orphan permanent file should be cleaned up: {p}"


@pytest.mark.asyncio
async def test_order_print_service_uses_cache(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    monkeypatch.setattr(settings, "BACKEND_INTERNAL_CALLBACK_URL", "http://app")
    monkeypatch.setattr(settings, "APP_PUBLIC_URL", "http://app")
    monkeypatch.setattr(settings, "ONLYOFFICE_PUBLIC_URL", "http://localhost:8085")
    monkeypatch.setattr(settings, "ONLYOFFICE_INTERNAL_URL", "http://onlyoffice:80")

    source_docx = tmp_path / "2026" / "order.docx"
    source_docx.parent.mkdir(parents=True, exist_ok=True)
    source_docx.write_bytes(b"PK")

    service = OrderPrintService()
    convert_calls: list[str] = []
    download_calls: list[str] = []

    async def fake_convert(entity_kind: str, entity_id: int, docx_path: Path, cache_key: str) -> str:
        assert entity_kind == "order"
        assert entity_id == 7
        convert_calls.append(cache_key)
        return "http://onlyoffice/cache/converted.pdf"

    async def fake_download(file_url: str) -> bytes:
        download_calls.append(file_url)
        return b"%PDF-1.7 test"

    monkeypatch.setattr(service, "_convert_docx_to_pdf", fake_convert)
    monkeypatch.setattr(service, "_download_pdf", fake_download)

    first = await service.get_or_create_pdf("order", 7, source_docx)
    second = await service.get_or_create_pdf("order", 7, source_docx)
    new_mtime = time.time() + 5
    os.utime(source_docx, (new_mtime, new_mtime))
    third = await service.get_or_create_pdf("order", 7, source_docx)

    assert first == second
    assert third != first
    assert not first.exists()
    assert third.exists()
    assert third.read_bytes() == b"%PDF-1.7 test"
    assert len(convert_calls) == 2
    assert len(download_calls) == 2


@pytest.mark.asyncio
async def test_order_print_service_uses_distinct_cache_namespace(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")
    monkeypatch.setattr(settings, "BACKEND_INTERNAL_CALLBACK_URL", "http://app")
    monkeypatch.setattr(settings, "APP_PUBLIC_URL", "http://app")
    monkeypatch.setattr(settings, "ONLYOFFICE_PUBLIC_URL", "http://localhost:8085")
    monkeypatch.setattr(settings, "ONLYOFFICE_INTERNAL_URL", "http://onlyoffice:80")

    source_docx = tmp_path / "2026" / "doc.docx"
    source_docx.parent.mkdir(parents=True, exist_ok=True)
    source_docx.write_bytes(b"PK")

    service = OrderPrintService()
    convert_calls: list[str] = []

    async def fake_convert(_entity_kind: str, _entity_id: int, _docx_path: Path, cache_key: str) -> str:
        convert_calls.append(cache_key)
        return f"http://onlyoffice/cache/{cache_key}.pdf"

    async def fake_download(file_url: str) -> bytes:
        return f"%PDF-{file_url}".encode("utf-8")

    monkeypatch.setattr(service, "_convert_docx_to_pdf", fake_convert)
    monkeypatch.setattr(service, "_download_pdf", fake_download)

    order_pdf = await service.get_or_create_pdf("order", 42, source_docx)
    notification_pdf = await service.get_or_create_pdf("notification", 42, source_docx)

    assert order_pdf != notification_pdf
    assert order_pdf.exists()
    assert notification_pdf.exists()
    assert order_pdf.name.startswith("order-42-")
    assert notification_pdf.name.startswith("notification-42-")
    assert len(convert_calls) == 2


@pytest.mark.asyncio
async def test_order_print_pdf_endpoint_returns_inline_file(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "ONLYOFFICE_ENABLED", True)
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))

    docx_path = tmp_path / "2026" / "order.docx"
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    docx_path.write_bytes(b"PK")

    pdf_path = tmp_path / ".print_cache" / "order-1-123.pdf"
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    pdf_path.write_bytes(b"%PDF-1.7")

    async def fake_get_by_id(_db, _order_id):
        return SimpleNamespace(file_path="2026/order.docx")

    async def fake_get_or_create_pdf(_entity_kind, _order_id, _docx_path: Path):
        return pdf_path

    monkeypatch.setattr(order_service, "get_by_id", fake_get_by_id)
    monkeypatch.setattr(order_print_service, "get_or_create_pdf", fake_get_or_create_pdf)

    response = await print_order_pdf(order_id=1, db=_db(), current_user="admin")

    assert response.media_type == "application/pdf"
    assert response.headers["content-disposition"].startswith("inline;")


@pytest.mark.asyncio
async def test_order_print_pdf_endpoint_404_when_file_missing(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "ONLYOFFICE_ENABLED", True)
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))

    async def fake_get_by_id(_db, _order_id):
        return SimpleNamespace(file_path="2026/missing.docx")

    monkeypatch.setattr(order_service, "get_by_id", fake_get_by_id)

    with pytest.raises(HRMSException) as exc_info:
        await print_order_pdf(order_id=1, db=_db(), current_user="admin")

    assert exc_info.value.status_code == 404
    assert exc_info.value.error_code == "order_file_missing"


@pytest.mark.asyncio
async def test_order_print_pdf_endpoint_propagates_conversion_error(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "ONLYOFFICE_ENABLED", True)
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))

    docx_path = tmp_path / "2026" / "order.docx"
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    docx_path.write_bytes(b"PK")

    async def fake_get_by_id(_db, _order_id):
        return SimpleNamespace(file_path="2026/order.docx")

    async def fake_get_or_create_pdf(_entity_kind, _order_id, _docx_path: Path):
        raise HRMSException("OnlyOffice conversion failed", "order_pdf_convert_failed", status_code=502)

    monkeypatch.setattr(order_service, "get_by_id", fake_get_by_id)
    monkeypatch.setattr(order_print_service, "get_or_create_pdf", fake_get_or_create_pdf)

    with pytest.raises(HRMSException) as exc_info:
        await print_order_pdf(order_id=1, db=_db(), current_user="admin")

    assert exc_info.value.status_code == 502
    assert exc_info.value.error_code == "order_pdf_convert_failed"


@pytest.mark.asyncio
async def test_get_current_user_or_onlyoffice_valid_onlyoffice_token(monkeypatch):
    from app.api.deps import get_current_user_or_onlyoffice

    monkeypatch.setattr(settings, "ONLYOFFICE_ENABLED", True)
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")

    token = jwt.encode({"status": 2}, "test-secret", algorithm="HS256")
    
    headers = {"Authorization": f"Bearer {token}"}
    mock_request = SimpleNamespace(headers=headers, method="GET", query_params={})

    result = await get_current_user_or_onlyoffice(request=cast(Request, mock_request), db=_db())
    assert result == "onlyoffice_server"


@pytest.mark.asyncio
async def test_get_current_user_or_onlyoffice_fallback_to_user(monkeypatch):
    from fastapi import HTTPException
    from app.api.deps import get_current_user_or_onlyoffice

    monkeypatch.setattr(settings, "ONLYOFFICE_ENABLED", True)
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")

    headers = {"Authorization": "Bearer invalid-token"}
    mock_request = SimpleNamespace(headers=headers, method="GET", query_params={})

    with pytest.raises(HTTPException) as exc_info:
        await get_current_user_or_onlyoffice(request=cast(Request, mock_request), db=_db())
    
    assert exc_info.value.status_code == 401


@pytest.mark.asyncio
async def test_get_current_user_or_onlyoffice_query_token(monkeypatch):
    from app.api.deps import get_current_user_or_onlyoffice
    from fastapi import HTTPException

    monkeypatch.setattr(settings, "ONLYOFFICE_ENABLED", True)
    monkeypatch.setattr(settings, "ONLYOFFICE_JWT_SECRET", "test-secret")

    mock_request = SimpleNamespace(
        headers={},
        method="GET",
        query_params={"token": "invalid-token"}
    )

    with pytest.raises(HTTPException) as exc_info:
        await get_current_user_or_onlyoffice(request=cast(Request, mock_request), db=_db())
    
    assert exc_info.value.status_code == 401


# === #30: Commit from server metadata ===


@pytest.mark.asyncio
async def test_commit_single_order_from_metadata(monkeypatch, tmp_path):
    """#30 AC1/AC6: commit creates order from server metadata without frontend payload."""
    import json
    from app.services.order_draft_service import order_draft_service
    from app.services.order_service import OrderService

    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    order_draft_service._drafts_dir = tmp_path / ".drafts"
    order_draft_service.ensure_drafts_dir()

    draft_id = "aaaaaaaa-1111-2222-3333-444444444444"
    # Create draft docx
    draft_path = order_draft_service._drafts_dir / f"{draft_id}_order.docx"
    draft_path.write_bytes(b"draft-content")

    # Create metadata
    metadata = {
        "draft_id": draft_id,
        "kind": "single_order",
        "order_type_code": "vacation_paid",
        "payload": {
            "employee_id": 5,
            "order_type_id": 2,
            "order_date": "2026-08-01",
            "order_number": "99-К",
            "notes": None,
            "extra_fields": {"vacation_start": "2026-08-10", "vacation_end": "2026-08-20"},
        },
        "created_by": "admin",
        "created_at": "2026-08-01T12:00:00+00:00",
        "status": "draft",
        "schema_version": 1,
    }
    order_draft_service.save_draft_metadata(draft_id, metadata)

    svc = OrderService()
    fake_order = SimpleNamespace(id=77, order_number="99-К")

    async def fake_create_order(db, data):
        # Verify OrderCreate was built from metadata
        assert data.employee_id == 5
        assert data.order_type_id == 2
        assert data.order_date == date(2026, 8, 1)
        assert data.order_number == "99-К"
        assert data.extra_fields == {"vacation_start": "2026-08-10", "vacation_end": "2026-08-20"}
        assert data.draft_id == draft_id
        return fake_order

    monkeypatch.setattr(svc, "create_order", fake_create_order)

    db = MagicMock()
    order = await svc.create_single_order_from_draft(db, draft_id)
    assert order.id == 77


@pytest.mark.asyncio
async def test_commit_draft_without_metadata_raises_outdated(monkeypatch, tmp_path):
    """#30 AC3: draft without metadata is rejected with 'Черновик устарел' message."""
    from app.services.order_draft_service import order_draft_service
    from app.services.order_service import OrderService

    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    order_draft_service._drafts_dir = tmp_path / ".drafts"
    order_draft_service.ensure_drafts_dir()

    draft_id = "bbbbbbbb-1111-2222-3333-444444444444"
    # Create draft docx but NO metadata
    draft_path = order_draft_service._drafts_dir / f"{draft_id}_order.docx"
    draft_path.write_bytes(b"draft-no-meta")

    svc = OrderService()
    db = MagicMock()

    with pytest.raises(HRMSException) as exc_info:
        await svc.create_single_order_from_draft(db, draft_id)

    assert exc_info.value.status_code == 409
    assert "Черновик устарел" in exc_info.value.message
    assert exc_info.value.error_code == "draft_outdated"


@pytest.mark.asyncio
async def test_duplicate_commit_returns_success(monkeypatch, tmp_path):
    """#30 AC5: duplicate commit of same draft doesn't create second order (silent success)."""
    from app.api.onlyoffice import commit_order_draft
    from app.services.order_draft_service import order_draft_service

    monkeypatch.setattr(settings, "ONLYOFFICE_ENABLED", True)
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    order_draft_service._drafts_dir = tmp_path / ".drafts"
    order_draft_service.ensure_drafts_dir()

    draft_id = "cccccccc-1111-2222-3333-444444444444"
    draft_path = order_draft_service._drafts_dir / f"{draft_id}_order.docx"
    draft_path.write_bytes(b"draft-dup")

    # Simulate already-claimed lock (first commit succeeded)
    lock_path = order_draft_service._commit_lock_path(draft_id)
    lock_path.write_bytes(b"1")

    db = MagicMock()
    result = await commit_order_draft(draft_id=draft_id, db=db, current_user="admin")

    assert result["duplicate"] is True
    assert "уже создан" in result["message"]


@pytest.mark.asyncio
async def test_failed_commit_releases_lock_for_retry(monkeypatch, tmp_path):
    """#30 AC4: on failed save, draft remains available for retry (lock released)."""
    from app.api.onlyoffice import commit_order_draft
    from app.services.order_draft_service import order_draft_service
    from app.services.order_service import order_service as order_service_singleton

    monkeypatch.setattr(settings, "ONLYOFFICE_ENABLED", True)
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    order_draft_service._drafts_dir = tmp_path / ".drafts"
    order_draft_service.ensure_drafts_dir()

    draft_id = "dddddddd-1111-2222-3333-444444444444"
    draft_path = order_draft_service._drafts_dir / f"{draft_id}_order.docx"
    draft_path.write_bytes(b"draft-retry")

    # Metadata exists
    metadata = {
        "draft_id": draft_id,
        "kind": "single_order",
        "order_type_code": "general_order",
        "payload": {
            "employee_id": None,
            "order_type_id": 1,
            "order_date": "2026-08-01",
            "order_number": "50",
            "notes": None,
            "extra_fields": None,
        },
        "created_by": "admin",
        "created_at": "2026-08-01T12:00:00+00:00",
        "status": "draft",
        "schema_version": 1,
    }
    order_draft_service.save_draft_metadata(draft_id, metadata)

    # Make create_single_order_from_draft fail
    async def failing_create(db, draft_id_arg):
        raise RuntimeError("DB connection lost")

    monkeypatch.setattr(order_service_singleton, "create_single_order_from_draft", failing_create)

    db = MagicMock()
    with pytest.raises(RuntimeError, match="DB connection lost"):
        await commit_order_draft(draft_id=draft_id, db=db, current_user="admin")

    # Lock must be released so retry is possible
    lock_path = order_draft_service._commit_lock_path(draft_id)
    assert not lock_path.exists(), "commit lock must be released after failure"
    # Draft file still exists
    assert draft_path.exists()


# --- Tests for list_drafts (#32) ---


def test_list_drafts_returns_metadata(monkeypatch, tmp_path):
    """list_drafts returns metadata for existing drafts sorted by created_at desc."""
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    svc = OrderDraftService()
    svc.ensure_drafts_dir()

    # Create two drafts with docx files and metadata
    draft_id_1 = "aaaaaaaa-1111-2222-3333-444444444444"
    draft_id_2 = "bbbbbbbb-1111-2222-3333-444444444444"

    (svc._drafts_dir / f"{draft_id_1}_order1.docx").write_bytes(b"PK1")
    (svc._drafts_dir / f"{draft_id_2}_order2.docx").write_bytes(b"PK2")

    meta1 = {
        "draft_id": draft_id_1,
        "kind": "single_order",
        "order_type_code": "vacation_paid",
        "payload": {"employee_id": 1, "order_number": "10-К", "order_date": "2026-07-01"},
        "created_by": "admin",
        "created_at": "2026-07-01T10:00:00+00:00",
        "status": "draft",
        "schema_version": 1,
    }
    meta2 = {
        "draft_id": draft_id_2,
        "kind": "group_order",
        "order_type_code": "vacation_unpaid_group",
        "payload": {"order_number": "20-К", "order_date": "2026-08-01", "employees": [{"employee_id": 1}, {"employee_id": 2}]},
        "created_by": "admin",
        "created_at": "2026-08-01T12:00:00+00:00",
        "status": "draft",
        "schema_version": 1,
    }
    svc.save_draft_metadata(draft_id_1, meta1)
    svc.save_draft_metadata(draft_id_2, meta2)

    result = svc.list_drafts()
    assert len(result) == 2
    # Newest first
    assert result[0]["draft_id"] == draft_id_2
    assert result[1]["draft_id"] == draft_id_1
    assert result[0]["kind"] == "group_order"
    assert result[1]["order_type_code"] == "vacation_paid"


def test_list_drafts_skips_orphaned_metadata(monkeypatch, tmp_path):
    """list_drafts skips metadata whose docx file is missing."""
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    svc = OrderDraftService()
    svc.ensure_drafts_dir()

    # Draft with docx
    draft_id_ok = "cccccccc-1111-2222-3333-444444444444"
    (svc._drafts_dir / f"{draft_id_ok}_order.docx").write_bytes(b"PK")
    svc.save_draft_metadata(draft_id_ok, {
        "draft_id": draft_id_ok,
        "kind": "single_order",
        "order_type_code": "general_order",
        "payload": {"order_number": "5"},
        "created_by": "admin",
        "created_at": "2026-07-15T09:00:00+00:00",
        "status": "draft",
        "schema_version": 1,
    })

    # Orphaned metadata (no docx)
    draft_id_orphan = "dddddddd-5555-6666-7777-888888888888"
    svc.save_draft_metadata(draft_id_orphan, {
        "draft_id": draft_id_orphan,
        "kind": "single_order",
        "order_type_code": "general_order",
        "payload": {"order_number": "6"},
        "created_by": "admin",
        "created_at": "2026-07-16T09:00:00+00:00",
        "status": "draft",
        "schema_version": 1,
    })

    result = svc.list_drafts()
    assert len(result) == 1
    assert result[0]["draft_id"] == draft_id_ok


def test_list_drafts_empty_when_no_drafts(monkeypatch, tmp_path):
    """list_drafts returns empty list when no drafts exist."""
    monkeypatch.setattr(settings, "ORDERS_PATH", str(tmp_path))
    svc = OrderDraftService()

    result = svc.list_drafts()
    assert result == []

